"""
Telco Threat Intelligence Report Generator — Web UI
Run: python app.py   (or via Generate_Report.bat)
Then open: http://localhost:5000
"""

import io
import os
import sys
import tempfile
import threading
import webbrowser
import datetime
import yaml

from flask import Flask, request, send_file, jsonify, render_template_string

# ── Try reportlab ──────────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, PageBreak, HRFlowable
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.colors import HexColor
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# ── Try python-docx ────────────────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ══════════════════════════════════════════════════════════════════════════════
#  PDF ENGINE
# ══════════════════════════════════════════════════════════════════════════════

C = {
    "bg":     "#0D1B2A", "mid":   "#1E2D3D", "card":  "#2A3D52",
    "teal":   "#00C2CB", "yel":   "#F4D35E", "red":   "#E63946",
    "orange": "#F7A400", "green": "#2ECC71", "white": "#FFFFFF",
    "light":  "#C8D6E5", "muted": "#7F8C8D",
}
MARGIN = 1.5 * 28.35

def hc(h): return HexColor(h) if REPORTLAB_OK else None
def S(name, **kw): return ParagraphStyle(name, **kw) if REPORTLAB_OK else None

ST = {}

def init_styles():
    global ST
    ST = {
        "title":  S("T",  fontName="Helvetica-Bold", fontSize=26, leading=32,
                    textColor=hc(C["white"]), spaceAfter=6),
        "sub":    S("Su", fontName="Helvetica", fontSize=12, leading=16,
                    textColor=hc(C["teal"]), spaceAfter=4),
        "date":   S("D",  fontName="Helvetica-Bold", fontSize=10,
                    textColor=hc(C["light"])),
        "sh":     S("SH", fontName="Helvetica-Bold", fontSize=16, leading=20,
                    textColor=hc(C["teal"]), spaceBefore=12, spaceAfter=7),
        "ssh":    S("SS", fontName="Helvetica-Bold", fontSize=12, leading=16,
                    textColor=hc(C["yel"]), spaceBefore=9, spaceAfter=4),
        "body":   S("B",  fontName="Helvetica", fontSize=9.5, leading=14,
                    textColor=hc(C["light"]), alignment=TA_JUSTIFY, spaceAfter=5),
        "bullet": S("Bu", fontName="Helvetica", fontSize=9.5, leading=13,
                    textColor=hc(C["light"]),
                    leftIndent=14, firstLineIndent=-10, spaceBefore=2, spaceAfter=2),
        "label":  S("L",  fontName="Helvetica-Bold", fontSize=8,
                    textColor=hc(C["teal"])),
        "small":  S("Sm", fontName="Helvetica", fontSize=8, leading=11,
                    textColor=hc(C["muted"])),
        "mono":   S("Mo", fontName="Courier", fontSize=8.5, leading=12,
                    textColor=hc(C["light"])),
        "rec":    S("Re", fontName="Helvetica", fontSize=9,
                    textColor=hc(C["yel"]), leading=13),
    }

def sev_col(s):
    return hc({"CRITICAL": C["red"], "HIGH": C["orange"],
               "MEDIUM": C["teal"], "LOW": C["green"]}.get(str(s).upper(), C["light"]))

def base_ts():
    return TableStyle([
        ("BACKGROUND",    (0,0), (-1, 0), hc(C["mid"])),
        ("TEXTCOLOR",     (0,0), (-1, 0), hc(C["teal"])),
        ("FONTNAME",      (0,0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0,0), (-1,-1), 8),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [hc(C["card"]), hc(C["bg"])]),
        ("TEXTCOLOR",     (0,1), (-1,-1), hc(C["light"])),
        ("FONTNAME",      (0,1), (-1,-1), "Helvetica"),
        ("ALIGN",         (0,0), (-1,-1), "LEFT"),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",    (0,0), (-1,-1), 5),
        ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ("LEFTPADDING",   (0,0), (-1,-1), 6),
        ("BOX",           (0,0), (-1,-1), 1, hc(C["teal"])),
        ("INNERGRID",     (0,0), (-1,-1), 0.3, hc(C["mid"])),
    ])


class DarkCanvas(rl_canvas.Canvas):
    def __init__(self, filename, report_title="", period="", tlp="AMBER", org="", **kw):
        super().__init__(filename, **kw)
        self._rt, self._period, self._tlp, self._org = report_title, period, tlp, org
        self._paint()

    def _paint(self):
        w, h = A4
        self.saveState()
        self.setFillColor(hc(C["bg"])); self.rect(0, 0, w, h, fill=1, stroke=0)
        self.setFillColor(hc(C["teal"])); self.rect(0, h-8*mm, w, 8*mm, fill=1, stroke=0)
        self.setFillColor(hc(C["bg"])); self.setFont("Helvetica-Bold", 8)
        self.drawString(MARGIN, h-5.5*mm, f"{self._rt.upper()}  |  {self._period}")
        self.drawRightString(w-MARGIN, h-5.5*mm, f"TLP:{self._tlp}")
        self.setFillColor(hc(C["mid"])); self.rect(0, 0, w, 10*mm, fill=1, stroke=0)
        self.setFillColor(hc(C["teal"])); self.setFont("Helvetica", 7.5)
        self.drawString(MARGIN, 3.5*mm, f"Classification: TLP:{self._tlp} — Not for public distribution.")
        self.setFillColor(hc(C["muted"]))
        self.drawRightString(w-MARGIN, 3.5*mm, f"Page {self._pageNumber}  |  {self._org}")
        self.setFillColor(hc(C["mid"])); self.rect(0, 10*mm, 3*mm, h-18*mm, fill=1, stroke=0)
        self.restoreState()

    def showPage(self):
        super().showPage(); self._paint()


def build_pdf(data: dict, output_path: str, progress_cb=None):
    init_styles()
    def prog(msg):
        if progress_cb: progress_cb(msg)

    title  = data.get("title",  "Telco Cyber Threat Intelligence Report")
    period = data.get("period", "")
    org    = data.get("org",    "")
    author = data.get("author", "")
    tlp    = data.get("tlp",    "AMBER")

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=MARGIN+5*mm, rightMargin=MARGIN,
        topMargin=MARGIN+8*mm, bottomMargin=MARGIN+10*mm,
        title=title, author=author)

    story = []

    prog("Building cover...")
    story.append(Spacer(1, 2.5*cm))
    for line in title.split("—"):
        story.append(Paragraph(line.strip().upper(), ST["title"]))
    story.append(Paragraph(data.get("subtitle", "Communications Service Providers"), ST["sub"]))
    story.append(Spacer(1, 3*mm))
    story.append(Paragraph(period, ST["date"]))
    story.append(Spacer(1, 8*mm))

    tlv = data.get("threat_level", "HIGH")
    tl_tbl = Table([[Paragraph("OVERALL THREAT LEVEL", ST["label"]),
                     Paragraph(tlv, S("TLV", fontName="Helvetica-Bold", fontSize=14,
                                       textColor=sev_col(tlv), alignment=TA_CENTER))]],
                   colWidths=[100*mm, 60*mm])
    tl_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), hc(C["card"])),
        ("BOX",           (0,0), (-1,-1), 1.5, hc(C["teal"])),
        ("LINEAFTER",     (0,0), (0,-1),  1,   hc(C["teal"])),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",    (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
        ("LEFTPADDING",   (0,0), (-1,-1), 10),
    ]))
    story.append(tl_tbl)
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(
        f"Prepared by: {author}  |  Organization: {org}  |  "
        f"Date: {data.get('report_date','')}  |  TLP:{tlp}", ST["small"]))
    story.append(PageBreak())

    prog("Building executive summary...")
    story.append(Paragraph("EXECUTIVE SUMMARY", ST["sh"]))
    story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
    exec_para = data.get("exec_summary", "").strip()
    if exec_para:
        story.append(Paragraph(exec_para, ST["body"]))
    top_risks = [r.strip() for r in data.get("top_risks", "").splitlines() if r.strip()]
    if top_risks:
        story.append(Paragraph("Top Risks:", ST["label"]))
        for r in top_risks:
            story.append(Paragraph(f"• {r}", ST["bullet"]))
        story.append(Spacer(1, 4*mm))
    kf_text = data.get("key_findings", "").strip()
    kf_rows = []
    for line in kf_text.splitlines():
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 2:
            kf_rows.append((parts + ["HIGH"])[:3])
    if kf_rows:
        tbl = Table([["KEY FINDING", "THREAT ACTOR / VECTOR", "SEVERITY"]] + kf_rows,
                    colWidths=[88*mm, 57*mm, 22*mm])
        ts = base_ts()
        for i, row in enumerate(kf_rows, 1):
            ts.add("TEXTCOLOR", (2,i), (2,i), sev_col(row[2] if len(row)>2 else "HIGH"))
            ts.add("FONTNAME",  (2,i), (2,i), "Helvetica-Bold")
        tbl.setStyle(ts); story.append(tbl)
    story.append(PageBreak())

    prog("Building threat landscape...")
    story.append(Paragraph("THREAT LANDSCAPE", ST["sh"]))
    story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=10))
    domains = [
        ("signaling",    "Telephony Signaling Networks (SS7 / Diameter / GTP)"),
        ("core_5g",      "Core Network & Radio Access Networks (5G / LTE)"),
        ("enterprise_it","Enterprise IT — BSS/OSS & Identity Systems"),
        ("fraud",        "Telecom Fraud — SIM Swap, IRSF & AI-Enabled Vishing"),
        ("supply_chain", "Third-Party & Supply Chain Risk"),
    ]
    for key, domain_title in domains:
        overview    = data.get(f"{key}_overview", "").strip()
        incidents   = data.get(f"{key}_incidents", "").strip()
        mitigations = data.get(f"{key}_mitigations", "").strip()
        if not (overview or incidents or mitigations): continue
        story.append(Paragraph(domain_title, ST["ssh"]))
        if overview:
            story.append(Paragraph(overview, ST["body"]))
        if incidents:
            inc_rows = []
            for line in incidents.splitlines():
                if "|" in line:
                    parts = [p.strip() for p in line.split("|")]
                    if len(parts) >= 3:
                        inc_rows.append([f"{parts[0]} — {parts[3] if len(parts)>3 else parts[0]}",
                                         parts[1], parts[2]])
            if inc_rows:
                tbl = Table([["RECENT INCIDENTS", "DATE", "SEVERITY"]] + inc_rows,
                            colWidths=[112*mm, 20*mm, 35*mm])
                ts = base_ts()
                for i, row in enumerate(inc_rows, 1):
                    if row[2]:
                        ts.add("TEXTCOLOR", (2,i), (2,i), sev_col(row[2]))
                        ts.add("FONTNAME",  (2,i), (2,i), "Helvetica-Bold")
                tbl.setStyle(ts); story.append(tbl); story.append(Spacer(1,4*mm))
        if mitigations:
            story.append(Paragraph(f"Mitigations — {domain_title}", ST["label"]))
            for m in mitigations.splitlines():
                if m.strip(): story.append(Paragraph(f"• {m.strip()}", ST["bullet"]))
        story.append(Spacer(1, 5*mm))
    story.append(PageBreak())

    prog("Building actors...")
    actors_text = data.get("actors", "").strip()
    if actors_text:
        story.append(Paragraph("THREAT ACTORS & TTPs", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        actor_rows = []
        for line in actors_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 6: parts.append("")
                actor_rows.append(parts[:6])
        if actor_rows:
            tbl = Table([["ACTOR","ORIGIN","TARGET","INITIAL ACCESS","KEY TTPs","MITRE IDs"]]
                        + actor_rows, colWidths=[22*mm,18*mm,25*mm,30*mm,47*mm,25*mm])
            ts = base_ts()
            state_kw = ["china","russia","iran","dprk","prc","state"]
            for i, row in enumerate(actor_rows, 1):
                origin = row[1].lower() if len(row)>1 else ""
                col = hc(C["red"]) if any(k in origin for k in state_kw) else hc(C["orange"])
                ts.add("TEXTCOLOR", (1,i),(1,i), col)
                ts.add("FONTNAME",  (1,i),(1,i), "Helvetica-Bold")
                ts.add("TEXTCOLOR", (5,i),(5,i), hc(C["teal"]))
            tbl.setStyle(ts); story.append(tbl)
        story.append(PageBreak())

    prog("Building vulnerabilities...")
    vulns_text = data.get("vulnerabilities", "").strip()
    if vulns_text:
        story.append(Paragraph("VULNERABILITIES & EXPLOIT ALERT", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        vuln_rows = []
        for line in vulns_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 5: parts.append("")
                vuln_rows.append(parts[:5])
        if vuln_rows:
            tbl = Table([["CVE ID","VENDOR / PRODUCT","DESCRIPTION","CVSS","PATCH PRIORITY"]]
                        + vuln_rows, colWidths=[25*mm,28*mm,80*mm,12*mm,22*mm])
            ts = base_ts()
            pri_map = {"IMMEDIATE":C["red"],"HIGH":C["orange"],"MEDIUM":C["teal"],"LOW":C["green"]}
            for i, row in enumerate(vuln_rows, 1):
                pri = row[4].upper() if len(row)>4 else ""
                if pri in pri_map:
                    ts.add("TEXTCOLOR",(4,i),(4,i), hc(pri_map[pri]))
                    ts.add("FONTNAME", (4,i),(4,i), "Helvetica-Bold")
            tbl.setStyle(ts); story.append(tbl)
        story.append(Spacer(1, 6*mm))

    prog("Building malware...")
    malware_text = data.get("malware", "").strip()
    if malware_text:
        story.append(Paragraph("EMERGING MALWARE NOTIFICATION", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        for line in malware_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 5: parts.append("")
                name, date, sev, desc, rec = parts[0],parts[1],parts[2],parts[3],parts[4]
                sc = sev_col(sev)
                hdr_tbl = Table([[
                    Paragraph(name, S("MH",fontName="Helvetica-Bold",fontSize=10,textColor=hc(C["yel"]))),
                    Paragraph(date, ST["small"]),
                    Paragraph(sev,  S("MS",fontName="Helvetica-Bold",fontSize=9,textColor=sc,alignment=TA_CENTER)),
                ]], colWidths=[100*mm,30*mm,37*mm])
                hdr_tbl.setStyle(TableStyle([
                    ("BACKGROUND",(0,0),(-1,-1),hc(C["mid"])),
                    ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                    ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
                    ("LEFTPADDING",(0,0),(-1,-1),8),("BOX",(0,0),(-1,-1),1,hc(C["teal"])),
                ]))
                story.append(hdr_tbl)
                if desc:
                    bt = Table([[Paragraph(f"<b>Description:</b> {desc}", ST["body"])]], colWidths=[167*mm])
                    bt.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),hc(C["card"])),
                        ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8),
                        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
                    story.append(bt)
                if rec:
                    rt = Table([[Paragraph(f"<b>Recommendation:</b> {rec}", ST["rec"])]], colWidths=[167*mm])
                    rt.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),HexColor("#1A2B1A")),
                        ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8),
                        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),8),
                        ("BOX",(0,0),(-1,-1),0.5,hc(C["green"]))]))
                    story.append(rt)
                story.append(Spacer(1, 4*mm))

    prog("Building breaches...")
    breaches_text = data.get("breaches", "").strip()
    if breaches_text:
        story.append(Paragraph("BREACH NOTIFICATIONS", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        breach_rows = []
        for line in breaches_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 6: parts.append("")
                breach_rows.append(parts[:6])
        if breach_rows:
            tbl = Table([["ORGANIZATION","DATE","SECTOR","RECORDS EXPOSED","VECTOR","STATUS"]]
                        + breach_rows, colWidths=[32*mm,15*mm,18*mm,35*mm,38*mm,22*mm])
            ts = base_ts()
            for i, row in enumerate(breach_rows, 1):
                status = row[5].lower() if len(row)>5 else ""
                if "investigation" in status or "alleged" in status:
                    ts.add("TEXTCOLOR",(5,i),(5,i),hc(C["orange"]))
                    ts.add("FONTNAME", (5,i),(5,i),"Helvetica-Bold")
            tbl.setStyle(ts); story.append(tbl)
        story.append(Spacer(1, 6*mm))

    prog("Building IOCs...")
    iocs_text = data.get("iocs", "").strip()
    if iocs_text:
        story.append(Paragraph("INDICATORS OF COMPROMISE (IOCs)", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        story.append(Paragraph("Defanged IOCs for blocking/alerting.", ST["small"]))
        story.append(Spacer(1, 3*mm))
        ioc_rows = []
        for line in iocs_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 5: parts.append("")
                ioc_rows.append([parts[0],Paragraph(parts[1],ST["mono"]),parts[2],parts[3],parts[4]])
        if ioc_rows:
            tbl = Table([["TYPE","VALUE","CAMPAIGN / MALWARE","CONFIDENCE","ACTION"]] + ioc_rows,
                        colWidths=[22*mm,42*mm,55*mm,18*mm,30*mm])
            ts = base_ts()
            action_map = {"BLOCK":C["red"],"BLOCK AT STP":C["red"],"ALERT":C["orange"],"MONITOR":C["teal"]}
            conf_map   = {"HIGH":C["green"],"MEDIUM":C["orange"],"LOW":C["muted"]}
            for i, row in enumerate(ioc_rows, 1):
                act = str(row[4]).upper(); conf = str(row[3]).upper()
                if act in action_map:
                    ts.add("TEXTCOLOR",(4,i),(4,i),hc(action_map[act]))
                    ts.add("FONTNAME", (4,i),(4,i),"Helvetica-Bold")
                if conf in conf_map:
                    ts.add("TEXTCOLOR",(3,i),(3,i),hc(conf_map[conf]))
            tbl.setStyle(ts); story.append(tbl)

    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=hc(C["muted"]), spaceAfter=5))
    story.append(Paragraph(
        f"CLASSIFICATION: TLP:{tlp} — This report may not be shared beyond the recipient's "
        "organization. Based on information available at time of publication. Provided AS-IS "
        "without warranty. All IOCs are defanged for safe distribution.", ST["small"]))

    prog("Rendering PDF...")

    class BoundCanvas(DarkCanvas):
        def __init__(self, filename, **kw):
            super().__init__(filename, report_title=title, period=period, tlp=tlp, org=org, **kw)

    doc.build(story, canvasmaker=BoundCanvas)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def build_docx(data: dict, output_path: str, progress_cb=None):
    def prog(msg):
        if progress_cb: progress_cb(msg)

    title  = data.get("title",  "Telco Cyber Threat Intelligence Report")
    period = data.get("period", "")
    org    = data.get("org",    "")
    author = data.get("author", "")
    tlp    = data.get("tlp",    "AMBER")
    tlv    = data.get("threat_level", "HIGH")

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

    def hex2rgb(h):
        h = h.lstrip("#")
        return RGBColor(int(h[0:2],16),int(h[2:4],16),int(h[4:6],16))

    def sev_rgb(s):
        return hex2rgb({"CRITICAL":"E63946","HIGH":"F7A400",
                        "MEDIUM":"00C2CB","LOW":"2ECC71"}.get(str(s).upper(),"C8D6E5"))

    def cell_bg(cell, hex_col):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"),hex_col.lstrip("#")); tcPr.append(shd)

    def add_heading(text, level=1):
        p=doc.add_paragraph()
        run=p.add_run(text.upper() if level==1 else text)
        run.font.size=Pt(16 if level==1 else 11); run.font.bold=True
        run.font.color.rgb=hex2rgb("0D1B2A" if level==1 else "00C2CB")
        p.space_before=Pt(14 if level==1 else 10); p.space_after=Pt(6)
        if level==1:
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
            bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
            bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"1")
            bot.set(qn("w:color"),"00C2CB"); pBdr.append(bot); pPr.append(pBdr)

    def add_bullet(text):
        p=doc.add_paragraph(style="List Bullet")
        run=p.add_run(text); run.font.size=Pt(9.5); p.space_after=Pt(2)

    def make_table(headers, rows, ncols):
        tbl=doc.add_table(rows=1+len(rows),cols=ncols); tbl.style="Table Grid"
        hrow=tbl.rows[0]
        for i,col in enumerate(headers):
            cell_bg(hrow.cells[i],"0D1B2A")
            rc=hrow.cells[i].paragraphs[0].add_run(col)
            rc.font.bold=True; rc.font.size=Pt(8); rc.font.color.rgb=hex2rgb("00C2CB")
        return tbl

    # Header / Footer
    prog("Building cover...")
    section=doc.sections[0]
    hp=section.header.paragraphs[0]
    hp.text=f"{title}  |  {period}  |  TLP:{tlp}"
    hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=hex2rgb("0D1B2A")
    fp=section.footer.paragraphs[0]
    fp.text=f"Classification: TLP:{tlp} — Not for public distribution.  |  {org}"
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in fp.runs: r.font.size=Pt(7); r.font.color.rgb=hex2rgb("7F8C8D")

    # Cover
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.space_before=Pt(72)
    run=p.add_run(title.upper()); run.font.size=Pt(22); run.font.bold=True
    run.font.color.rgb=hex2rgb("0D1B2A")
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run2=p2.add_run(data.get("subtitle","Communications Service Providers"))
    run2.font.size=Pt(13); run2.font.color.rgb=hex2rgb("00C2CB")
    doc.add_paragraph()
    tl_tbl=doc.add_table(rows=1,cols=2); tl_tbl.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cell_bg(tl_tbl.cell(0,0),"0D1B2A"); cell_bg(tl_tbl.cell(0,1),"0D1B2A")
    r0=tl_tbl.cell(0,0).paragraphs[0].add_run("OVERALL THREAT LEVEL")
    r0.font.size=Pt(9); r0.font.bold=True; r0.font.color.rgb=hex2rgb("00C2CB")
    tl_tbl.cell(0,1).paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    r1=tl_tbl.cell(0,1).paragraphs[0].add_run(tlv)
    r1.font.size=Pt(16); r1.font.bold=True; r1.font.color.rgb=sev_rgb(tlv)
    doc.add_paragraph()
    meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER
    mr=meta.add_run(f"Prepared by: {author}  |  Organization: {org}  |  "
                    f"Date: {data.get('report_date','')}  |  TLP:{tlp}")
    mr.font.size=Pt(9); mr.font.color.rgb=hex2rgb("7F8C8D")
    doc.add_page_break()

    # Exec summary
    prog("Building executive summary...")
    add_heading("Executive Summary")
    exec_para=data.get("exec_summary","").strip()
    if exec_para:
        ep=doc.add_paragraph(exec_para)
        if ep.runs: ep.runs[0].font.size=Pt(10)
    top_risks=[r.strip() for r in data.get("top_risks","").splitlines() if r.strip()]
    if top_risks:
        lp=doc.add_paragraph(); lr=lp.add_run("Top Risks:")
        lr.font.bold=True; lr.font.size=Pt(10); lr.font.color.rgb=hex2rgb("00C2CB")
        for r in top_risks: add_bullet(r)
    kf_text=data.get("key_findings","").strip(); kf_rows=[]
    for line in kf_text.splitlines():
        parts=[pp.strip() for pp in line.split("|")]
        if len(parts)>=2: kf_rows.append((parts+["HIGH"])[:3])
    if kf_rows:
        doc.add_paragraph()
        tbl=make_table(["KEY FINDING","ACTOR / VECTOR","SEVERITY"],kf_rows,3)
        for ri,row_data in enumerate(kf_rows,1):
            cells=tbl.rows[ri].cells; bg="F5F7FA" if ri%2==0 else "FFFFFF"
            for ci,text in enumerate(row_data):
                cell_bg(cells[ci],bg); cells[ci].text=text
                if ci==2 and text:
                    for para in cells[ci].paragraphs:
                        for run in para.runs: run.font.bold=True; run.font.color.rgb=sev_rgb(text)
    doc.add_page_break()

    # Threat landscape
    prog("Building threat landscape...")
    add_heading("Threat Landscape")
    domains=[("signaling","Telephony Signaling Networks (SS7 / Diameter / GTP)"),
             ("core_5g","Core Network & Radio Access Networks (5G / LTE)"),
             ("enterprise_it","Enterprise IT — BSS/OSS & Identity Systems"),
             ("fraud","Telecom Fraud — SIM Swap, IRSF & AI-Enabled Vishing"),
             ("supply_chain","Third-Party & Supply Chain Risk")]
    for key,domain_title in domains:
        overview=data.get(f"{key}_overview","").strip()
        incidents=data.get(f"{key}_incidents","").strip()
        mitigations=data.get(f"{key}_mitigations","").strip()
        if not (overview or incidents or mitigations): continue
        add_heading(domain_title,level=2)
        if overview:
            ep=doc.add_paragraph(overview)
            if ep.runs: ep.runs[0].font.size=Pt(10)
        if incidents:
            inc_rows=[]
            for line in incidents.splitlines():
                if "|" in line:
                    parts=[pp.strip() for pp in line.split("|")]
                    while len(parts)<4: parts.append("")
                    inc_rows.append(parts[:4])
            if inc_rows:
                tbl=make_table(["INCIDENT","DATE","SEVERITY","DESCRIPTION"],inc_rows,4)
                for ri,parts in enumerate(inc_rows,1):
                    cells=tbl.rows[ri].cells; bg="F5F7FA" if ri%2==0 else "FFFFFF"
                    for ci,text in enumerate(parts):
                        cell_bg(cells[ci],bg); cells[ci].text=text
                        if ci==2 and text:
                            for para in cells[ci].paragraphs:
                                for run in para.runs: run.font.bold=True; run.font.color.rgb=sev_rgb(text)
                doc.add_paragraph()
        if mitigations:
            lp=doc.add_paragraph(); lr=lp.add_run("Mitigations")
            lr.font.bold=True; lr.font.size=Pt(9); lr.font.color.rgb=hex2rgb("2ECC71")
            for m in mitigations.splitlines():
                if m.strip(): add_bullet(m.strip())
    doc.add_page_break()

    # Actors
    actors_text=data.get("actors","").strip()
    if actors_text:
        prog("Building actors...")
        add_heading("Threat Actors & TTPs"); actor_rows=[]
        for line in actors_text.splitlines():
            if "|" in line:
                parts=[pp.strip() for pp in line.split("|")]
                while len(parts)<6: parts.append("")
                actor_rows.append(parts[:6])
        if actor_rows:
            tbl=make_table(["ACTOR","ORIGIN","TARGET","INITIAL ACCESS","KEY TTPs","MITRE IDs"],actor_rows,6)
            state_kw=["china","russia","iran","dprk","prc","state"]
            for ri,row_data in enumerate(actor_rows,1):
                cells=tbl.rows[ri].cells; bg="F5F7FA" if ri%2==0 else "FFFFFF"
                for ci,text in enumerate(row_data):
                    cell_bg(cells[ci],bg); cells[ci].text=text
                    if ci==1:
                        col_rgb=hex2rgb("E63946") if any(k in text.lower() for k in state_kw) else hex2rgb("F7A400")
                        for para in cells[ci].paragraphs:
                            for run in para.runs: run.font.bold=True; run.font.color.rgb=col_rgb
                    if ci==5:
                        for para in cells[ci].paragraphs:
                            for run in para.runs: run.font.color.rgb=hex2rgb("00C2CB")

    # Vulns
    vulns_text=data.get("vulnerabilities","").strip()
    if vulns_text:
        prog("Building vulnerabilities...")
        add_heading("Vulnerabilities & Exploit Alert"); vuln_rows=[]
        for line in vulns_text.splitlines():
            if "|" in line:
                parts=[pp.strip() for pp in line.split("|")]
                while len(parts)<5: parts.append("")
                vuln_rows.append(parts[:5])
        if vuln_rows:
            tbl=make_table(["CVE ID","VENDOR / PRODUCT","DESCRIPTION","CVSS","PATCH PRIORITY"],vuln_rows,5)
            pri_map={"IMMEDIATE":"E63946","HIGH":"F7A400","MEDIUM":"00C2CB","LOW":"2ECC71"}
            for ri,row_data in enumerate(vuln_rows,1):
                cells=tbl.rows[ri].cells; bg="F5F7FA" if ri%2==0 else "FFFFFF"
                for ci,text in enumerate(row_data):
                    cell_bg(cells[ci],bg); cells[ci].text=text
                    if ci==4 and text.upper() in pri_map:
                        for para in cells[ci].paragraphs:
                            for run in para.runs: run.font.bold=True; run.font.color.rgb=hex2rgb(pri_map[text.upper()])

    # Malware
    malware_text=data.get("malware","").strip()
    if malware_text:
        prog("Building malware...")
        add_heading("Emerging Malware Notification")
        for line in malware_text.splitlines():
            if "|" in line:
                parts=[pp.strip() for pp in line.split("|")]
                while len(parts)<5: parts.append("")
                name,date,sev,desc,rec=parts[0],parts[1],parts[2],parts[3],parts[4]
                hp=doc.add_paragraph(); hr=hp.add_run(f"[{sev}]  {name}  —  {date}")
                hr.font.bold=True; hr.font.size=Pt(11); hr.font.color.rgb=sev_rgb(sev)
                if desc:
                    dp=doc.add_paragraph()
                    dr=dp.add_run("Description: "); dr.font.bold=True; dr.font.size=Pt(9.5)
                    dr2=dp.add_run(desc); dr2.font.size=Pt(9.5)
                if rec:
                    rp=doc.add_paragraph()
                    rr=rp.add_run("Recommendation: "); rr.font.bold=True; rr.font.size=Pt(9.5)
                    rr.font.color.rgb=hex2rgb("2ECC71")
                    rr2=rp.add_run(rec); rr2.font.size=Pt(9.5)
                doc.add_paragraph()

    # Breaches
    breaches_text=data.get("breaches","").strip()
    if breaches_text:
        prog("Building breaches...")
        add_heading("Breach Notifications"); breach_rows=[]
        for line in breaches_text.splitlines():
            if "|" in line:
                parts=[pp.strip() for pp in line.split("|")]
                while len(parts)<6: parts.append("")
                breach_rows.append(parts[:6])
        if breach_rows:
            tbl=make_table(["ORGANIZATION","DATE","SECTOR","RECORDS EXPOSED","VECTOR","STATUS"],breach_rows,6)
            for ri,row_data in enumerate(breach_rows,1):
                cells=tbl.rows[ri].cells; bg="F5F7FA" if ri%2==0 else "FFFFFF"
                for ci,text in enumerate(row_data):
                    cell_bg(cells[ci],bg); cells[ci].text=text
                    if ci==5:
                        status=text.lower()
                        if "investigation" in status or "alleged" in status:
                            for para in cells[ci].paragraphs:
                                for run in para.runs: run.font.bold=True; run.font.color.rgb=hex2rgb("F7A400")

    # IOCs
    iocs_text=data.get("iocs","").strip()
    if iocs_text:
        prog("Building IOCs...")
        add_heading("Indicators of Compromise (IOCs)")
        np2=doc.add_paragraph("Defanged IOCs for blocking/alerting. All IOCs are safe to distribute.")
        if np2.runs: np2.runs[0].font.size=Pt(9); np2.runs[0].font.italic=True
        ioc_rows=[]
        for line in iocs_text.splitlines():
            if "|" in line:
                parts=[pp.strip() for pp in line.split("|")]
                while len(parts)<5: parts.append("")
                ioc_rows.append(parts[:5])
        if ioc_rows:
            tbl=make_table(["TYPE","VALUE","CAMPAIGN / MALWARE","CONFIDENCE","ACTION"],ioc_rows,5)
            action_map={"BLOCK":"E63946","BLOCK AT STP":"E63946","ALERT":"F7A400","MONITOR":"00C2CB"}
            conf_map={"HIGH":"2ECC71","MEDIUM":"F7A400","LOW":"7F8C8D"}
            for ri,row_data in enumerate(ioc_rows,1):
                cells=tbl.rows[ri].cells; bg="F5F7FA" if ri%2==0 else "FFFFFF"
                for ci,text in enumerate(row_data):
                    cell_bg(cells[ci],bg); cells[ci].text=text
                    if ci==1:
                        for para in cells[ci].paragraphs:
                            for run in para.runs: run.font.name="Courier New"; run.font.size=Pt(8.5)
                    if ci==3 and text.upper() in conf_map:
                        for para in cells[ci].paragraphs:
                            for run in para.runs: run.font.bold=True; run.font.color.rgb=hex2rgb(conf_map[text.upper()])
                    if ci==4 and text.upper() in action_map:
                        for para in cells[ci].paragraphs:
                            for run in para.runs: run.font.bold=True; run.font.color.rgb=hex2rgb(action_map[text.upper()])

    doc.add_paragraph()
    dp=doc.add_paragraph()
    dr=dp.add_run(f"CLASSIFICATION: TLP:{tlp} — This report may not be shared beyond the recipient's "
                  "organization. Based on information available at time of publication.")
    dr.font.size=Pt(8); dr.font.italic=True; dr.font.color.rgb=hex2rgb("7F8C8D")

    prog("Saving DOCX...")
    doc.save(output_path)
    return output_path


# ══════════════════════════════════════════════════════════════════════════════
#  FLASK WEB APP
# ══════════════════════════════════════════════════════════════════════════════

app = Flask(__name__)

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Telco Threat Intel Report Generator</title>
<style>
:root{--bg:#1a1a2e;--bg2:#16213e;--bg3:#0f3460;--teal:#00C2CB;--yel:#F4D35E;
      --red:#E63946;--white:#e0e0e0;--grey:#7f8c8d;--card:#2A3D52}
*{box-sizing:border-box;margin:0;padding:0}
body{background:var(--bg);color:var(--white);font-family:'Segoe UI',Tahoma,sans-serif;min-height:100vh}
header{background:var(--bg3);padding:14px 20px;display:flex;align-items:center;justify-content:space-between}
header h1{color:var(--teal);font-size:1.1rem;font-weight:700;letter-spacing:.5px}
header span{color:var(--yel);font-size:.8rem}
.tabs{display:flex;background:var(--bg2);border-bottom:2px solid var(--bg3);overflow-x:auto}
.tab{background:none;border:none;color:var(--white);padding:10px 18px;cursor:pointer;
     font-family:inherit;font-size:.85rem;white-space:nowrap;border-bottom:3px solid transparent;
     transition:.2s}
.tab:hover{background:var(--bg3)}
.tab.active{border-bottom-color:var(--teal);color:var(--teal);font-weight:600}
.content{display:none;padding:16px 20px;max-width:1100px;margin:0 auto}
.content.active{display:block}
.section{background:var(--bg2);border:1px solid var(--bg3);border-radius:6px;
         margin-bottom:14px;padding:14px}
.section h3{color:var(--teal);font-size:.85rem;margin-bottom:10px;text-transform:uppercase;
            letter-spacing:.5px}
.row{display:flex;align-items:center;gap:10px;margin-bottom:8px;flex-wrap:wrap}
.row label{min-width:160px;font-size:.83rem;color:var(--grey)}
.row input,.row select{background:var(--bg);border:1px solid var(--bg3);color:var(--white);
  padding:6px 10px;border-radius:4px;font-family:inherit;font-size:.83rem;flex:1;min-width:200px}
.row select option{background:var(--bg)}
.row input:focus,.row select:focus{outline:none;border-color:var(--teal)}
.lf{background:var(--bg2);border:1px solid var(--bg3);border-radius:5px;
    margin-bottom:12px;padding:12px}
.lf .lf-title{color:var(--teal);font-size:.8rem;font-weight:700;text-transform:uppercase;
              margin-bottom:8px;letter-spacing:.4px}
.hint{color:var(--grey);font-size:.75rem;margin-bottom:6px;font-style:italic;
      background:var(--bg3);padding:6px 8px;border-radius:3px;white-space:pre-wrap}
textarea{width:100%;background:var(--bg);border:1px solid var(--bg3);color:var(--white);
  padding:8px;border-radius:4px;font-family:'Consolas','Courier New',monospace;font-size:.8rem;
  resize:vertical;min-height:80px}
textarea:focus{outline:none;border-color:var(--teal)}
.domain{background:var(--bg);border:1px solid var(--bg3);border-radius:5px;
        margin-bottom:10px;padding:10px}
.domain h4{color:var(--yel);font-size:.8rem;margin-bottom:8px}
.domain label{color:var(--grey);font-size:.75rem;display:block;margin-top:6px;margin-bottom:3px}
footer{position:sticky;bottom:0;background:var(--bg2);border-top:1px solid var(--bg3);
       padding:10px 20px;display:flex;align-items:center;gap:10px;flex-wrap:wrap}
footer .status{color:var(--grey);font-size:.8rem;flex:1}
.btn{border:none;padding:9px 20px;border-radius:5px;font-family:inherit;font-size:.85rem;
     font-weight:700;cursor:pointer;transition:.2s}
.btn-pdf{background:var(--teal);color:#0D1B2A}
.btn-pdf:hover{background:#00a8b0}
.btn-docx{background:#2B579A;color:#fff}
.btn-docx:hover{background:#1e3f70}
.btn-sec{background:var(--bg3);color:var(--white)}
.btn-sec:hover{background:var(--bg)}
.btn:disabled{opacity:.5;cursor:not-allowed}
.ioc-note{background:rgba(244,211,94,.1);border:1px solid var(--yel);color:var(--yel);
           padding:8px 12px;border-radius:4px;font-size:.8rem;margin-bottom:10px}
input[type=file]{display:none}
</style>
</head>
<body>
<header>
  <h1>TELCO THREAT INTEL REPORT GENERATOR</h1>
  <span>v1.1  |  TLP:AMBER</span>
</header>

<div class="tabs">
  <button class="tab active" onclick="showTab('info')">Report Info</button>
  <button class="tab" onclick="showTab('exec')">Executive Summary</button>
  <button class="tab" onclick="showTab('landscape')">Threat Landscape</button>
  <button class="tab" onclick="showTab('actors')">Actors &amp; CVEs</button>
  <button class="tab" onclick="showTab('malware')">Malware &amp; Breaches</button>
  <button class="tab" onclick="showTab('iocs')">IOCs</button>
</div>

<!-- TAB 1: Report Info -->
<div class="content active" id="tab-info">
  <div class="section">
    <h3>Report Metadata</h3>
    <div class="row"><label>Report Title</label>
      <input id="title" value="Telco Cyber Threat Intelligence Report"></div>
    <div class="row"><label>Subtitle</label>
      <input id="subtitle" value="Communications Service Providers — India &amp; Global"></div>
    <div class="row"><label>Period Covered</label>
      <input id="period" value="April 2026"></div>
    <div class="row"><label>Report Date</label>
      <input id="report_date" type="date"></div>
    <div class="row"><label>Author</label>
      <input id="author" placeholder="e.g. Vardhan Chauhan"></div>
    <div class="row"><label>Organization</label>
      <input id="org" placeholder="e.g. Acme Telecom CSIRT"></div>
  </div>
  <div class="section">
    <h3>Classification</h3>
    <div class="row"><label>TLP Level</label>
      <select id="tlp"><option>WHITE</option><option>GREEN</option>
        <option selected>AMBER</option><option>RED</option></select></div>
    <div class="row"><label>Overall Threat Level</label>
      <select id="threat_level"><option>CRITICAL</option><option selected>HIGH</option>
        <option>MEDIUM</option><option>LOW</option></select></div>
  </div>
</div>

<!-- TAB 2: Executive Summary -->
<div class="content" id="tab-exec">
  <div class="lf"><div class="lf-title">Executive Summary Paragraph</div>
    <div class="hint">Write 3-5 sentences summarising the overall threat assessment for the period.</div>
    <textarea id="exec_summary" rows="5" placeholder="The threat landscape for telcos during this period was dominated by..."></textarea>
  </div>
  <div class="lf"><div class="lf-title">Top Risks (one per line)</div>
    <textarea id="top_risks" rows="5" placeholder="SS7 OTP interception targeting South Asian operators&#10;Salt Typhoon persistent access in Tier-1 backbone&#10;AI-enhanced vishing campaigns against telco helpdesks"></textarea>
  </div>
  <div class="lf"><div class="lf-title">Key Findings Table</div>
    <div class="hint">Format: Finding | Actor/Vector | Severity&#10;Severity: CRITICAL | HIGH | MEDIUM | LOW</div>
    <textarea id="key_findings" rows="6" placeholder="SS7 OTP interception active across South Asian carriers | Unattributed eCrime | CRITICAL&#10;Salt Typhoon persistence in Tier-1 backbone | China-nexus APT (PRC State) | HIGH"></textarea>
  </div>
</div>

<!-- TAB 3: Threat Landscape -->
<div class="content" id="tab-landscape">
  <div class="domain">
    <h4>SS7 / Diameter / GTP Signaling</h4>
    <label>Overview</label><textarea id="signaling_overview" rows="3"></textarea>
    <label>Incidents (Title | Date | Severity | Description)</label>
    <div class="hint">Example: SS7 OTP Interception | Mar 2026 | CRITICAL | Active campaign targeting Indian subs</div>
    <textarea id="signaling_incidents" rows="3"></textarea>
    <label>Mitigations (one per line)</label><textarea id="signaling_mitigations" rows="3"></textarea>
  </div>
  <div class="domain">
    <h4>Core Network &amp; 5G / LTE</h4>
    <label>Overview</label><textarea id="core_5g_overview" rows="3"></textarea>
    <label>Incidents (Title | Date | Severity | Description)</label>
    <textarea id="core_5g_incidents" rows="3"></textarea>
    <label>Mitigations</label><textarea id="core_5g_mitigations" rows="3"></textarea>
  </div>
  <div class="domain">
    <h4>Enterprise IT (BSS/OSS)</h4>
    <label>Overview</label><textarea id="enterprise_it_overview" rows="3"></textarea>
    <label>Incidents (Title | Date | Severity | Description)</label>
    <textarea id="enterprise_it_incidents" rows="3"></textarea>
    <label>Mitigations</label><textarea id="enterprise_it_mitigations" rows="3"></textarea>
  </div>
  <div class="domain">
    <h4>Telecom Fraud (SIM Swap / IRSF / Vishing)</h4>
    <label>Overview</label><textarea id="fraud_overview" rows="3"></textarea>
    <label>Incidents (Title | Date | Severity | Description)</label>
    <textarea id="fraud_incidents" rows="3"></textarea>
    <label>Mitigations</label><textarea id="fraud_mitigations" rows="3"></textarea>
  </div>
  <div class="domain">
    <h4>Supply Chain &amp; Third-Party</h4>
    <label>Overview</label><textarea id="supply_chain_overview" rows="3"></textarea>
    <label>Incidents (Title | Date | Severity | Description)</label>
    <textarea id="supply_chain_incidents" rows="3"></textarea>
    <label>Mitigations</label><textarea id="supply_chain_mitigations" rows="3"></textarea>
  </div>
</div>

<!-- TAB 4: Actors & CVEs -->
<div class="content" id="tab-actors">
  <div class="lf"><div class="lf-title">Threat Actors</div>
    <div class="hint">Format: Name | Origin | Target | Initial Access | TTPs | MITRE IDs&#10;Example: Salt Typhoon | China (PRC) | Tier-1 backbone | Exploit IOS-XE | LotL, GRE tunnel | T1190,T1133</div>
    <textarea id="actors" rows="8"></textarea>
  </div>
  <div class="lf"><div class="lf-title">Vulnerabilities</div>
    <div class="hint">Format: CVE ID | Vendor/Product | Description | CVSS | Patch Priority&#10;Patch Priority: IMMEDIATE | HIGH | MEDIUM | LOW&#10;Example: CVE-2026-2329 | Grandstream GXP1600 | Unauthenticated RCE via web API | 9.8 | IMMEDIATE</div>
    <textarea id="vulnerabilities" rows="8"></textarea>
  </div>
</div>

<!-- TAB 5: Malware & Breaches -->
<div class="content" id="tab-malware">
  <div class="lf"><div class="lf-title">Emerging Malware</div>
    <div class="hint">Format: Name | Date | Severity | Description | Recommendation&#10;Severity: CRITICAL | HIGH | MEDIUM&#10;Example: BPFDoor | Mar 2026 | HIGH | Passive BPF backdoor targeting Linux BSS/OSS | Audit raw socket listeners with ss -lp</div>
    <textarea id="malware" rows="8"></textarea>
  </div>
  <div class="lf"><div class="lf-title">Breach Notifications</div>
    <div class="hint">Format: Organization | Date | Sector | Records Exposed | Vector | Status&#10;Status: Confirmed | Under Investigation | Alleged | Contained&#10;Example: SK Telecom | Apr 2025 | Telco | 27M subscriber records | BPFDoor/LotL | Confirmed</div>
    <textarea id="breaches" rows="8"></textarea>
  </div>
</div>

<!-- TAB 6: IOCs -->
<div class="content" id="tab-iocs">
  <div class="ioc-note">Always defang IOCs: replace <b>.</b> with <b>[.]</b> in IPs and domains<br>
    Example: 185.220.101[.]47  or  malware-domain[.]com</div>
  <div class="lf"><div class="lf-title">Indicators of Compromise</div>
    <div class="hint">Format: Type | Value | Campaign/Malware | Confidence | Action&#10;Confidence: HIGH | MEDIUM | LOW&#10;Action: BLOCK | ALERT | MONITOR | BLOCK at STP&#10;Types: IPv4 | IPv6 | Domain | URL | File Hash (SHA256) | SCCP GT | MSISDN Range&#10;&#10;Example:&#10;IPv4 | 185.220.101[.]47 | GhostSignal SS7 C2 | HIGH | BLOCK&#10;Domain | update-secure[.]xyz | OysterLoader | HIGH | BLOCK&#10;SCCP GT | +442012345XXX | Rogue HLR | HIGH | BLOCK at STP</div>
    <textarea id="iocs" rows="15"></textarea>
  </div>
</div>

<footer>
  <span class="status" id="status">Ready</span>
  <button class="btn btn-sec" onclick="saveYaml()">Save YAML</button>
  <label class="btn btn-sec" style="cursor:pointer">Load YAML
    <input type="file" id="yaml_file" accept=".yaml,.yml" onchange="loadYaml(this)"></label>
  <button class="btn btn-docx" id="btn_docx" onclick="generate('docx')">GENERATE DOCX</button>
  <button class="btn btn-pdf"  id="btn_pdf"  onclick="generate('pdf')">GENERATE PDF</button>
</footer>

<script>
// ── Set today's date ──────────────────────────────────────────────────────────
document.getElementById('report_date').value = new Date().toISOString().slice(0,10);

// ── Tab switching ─────────────────────────────────────────────────────────────
function showTab(id){
  document.querySelectorAll('.content').forEach(e=>e.classList.remove('active'));
  document.querySelectorAll('.tab').forEach(e=>e.classList.remove('active'));
  document.getElementById('tab-'+id).classList.add('active');
  event.target.classList.add('active');
}

// ── Collect all form data ─────────────────────────────────────────────────────
function getAllData(){
  const ids=['title','subtitle','period','report_date','author','org',
             'tlp','threat_level','exec_summary','top_risks','key_findings',
             'signaling_overview','signaling_incidents','signaling_mitigations',
             'core_5g_overview','core_5g_incidents','core_5g_mitigations',
             'enterprise_it_overview','enterprise_it_incidents','enterprise_it_mitigations',
             'fraud_overview','fraud_incidents','fraud_mitigations',
             'supply_chain_overview','supply_chain_incidents','supply_chain_mitigations',
             'actors','vulnerabilities','malware','breaches','iocs'];
  const d={};
  ids.forEach(id=>{const el=document.getElementById(id); if(el) d[id]=el.value;});
  return d;
}

// ── Populate form from data object ───────────────────────────────────────────
function setAllData(d){
  Object.entries(d).forEach(([id,val])=>{
    const el=document.getElementById(id);
    if(el) el.value=val||'';
  });
}

// ── Generate PDF or DOCX ──────────────────────────────────────────────────────
async function generate(fmt){
  const btn=document.getElementById('btn_'+fmt);
  const status=document.getElementById('status');
  btn.disabled=true;
  status.textContent='Generating '+fmt.toUpperCase()+'...';
  try{
    const resp=await fetch('/generate_'+fmt,{
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify(getAllData())
    });
    if(!resp.ok){
      const err=await resp.json();
      alert('Error: '+(err.error||resp.statusText));
      return;
    }
    const blob=await resp.blob();
    const title=document.getElementById('title').value||'Telco_TI_Report';
    const safe=title.replace(/[^a-zA-Z0-9_\- ]/g,'').replace(/ /g,'_').slice(0,40);
    const ext=fmt==='pdf'?'.pdf':'.docx';
    const url=URL.createObjectURL(blob);
    const a=document.createElement('a');
    a.href=url; a.download=safe+ext; a.click();
    URL.revokeObjectURL(url);
    status.textContent=fmt.toUpperCase()+' saved to Downloads!';
    setTimeout(()=>status.textContent='Ready',4000);
  }catch(e){
    alert('Failed: '+e.message);
    status.textContent='Error';
  }finally{
    btn.disabled=false;
  }
}

// ── Save YAML ─────────────────────────────────────────────────────────────────
async function saveYaml(){
  const status=document.getElementById('status');
  status.textContent='Saving YAML...';
  try{
    const resp=await fetch('/save_yaml',{
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify(getAllData())
    });
    const blob=await resp.blob();
    const url=URL.createObjectURL(blob);
    const a=document.createElement('a'); a.href=url;
    a.download='report_manifest.yaml'; a.click();
    URL.revokeObjectURL(url);
    status.textContent='YAML saved!';
    setTimeout(()=>status.textContent='Ready',3000);
  }catch(e){alert('Save failed: '+e.message);}
}

// ── Load YAML ─────────────────────────────────────────────────────────────────
async function loadYaml(input){
  const file=input.files[0]; if(!file) return;
  const formData=new FormData(); formData.append('file',file);
  const status=document.getElementById('status');
  status.textContent='Loading YAML...';
  try{
    const resp=await fetch('/load_yaml',{method:'POST',body:formData});
    const data=await resp.json();
    if(data.error){alert('Error: '+data.error);return;}
    setAllData(data);
    status.textContent='Manifest loaded!';
    setTimeout(()=>status.textContent='Ready',3000);
  }catch(e){alert('Load failed: '+e.message);}
  input.value='';
}
</script>
</body>
</html>"""


@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/generate_pdf", methods=["POST"])
def generate_pdf_route():
    if not REPORTLAB_OK:
        return jsonify({"error": "reportlab not installed. Run: pip install reportlab"}), 500
    data = request.json
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        build_pdf(data, tmp.name)
        return send_file(tmp.name, mimetype="application/pdf",
                         as_attachment=True, download_name="Telco_TI_Report.pdf")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/generate_docx", methods=["POST"])
def generate_docx_route():
    if not DOCX_OK:
        return jsonify({"error": "python-docx not installed. Run: pip install python-docx"}), 500
    data = request.json
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp.close()
    try:
        build_docx(data, tmp.name)
        return send_file(tmp.name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name="Telco_TI_Report.docx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/save_yaml", methods=["POST"])
def save_yaml_route():
    data = request.json
    yaml_bytes = yaml.dump(data, default_flow_style=False, allow_unicode=True).encode()
    return send_file(io.BytesIO(yaml_bytes), mimetype="text/yaml",
                     as_attachment=True, download_name="report_manifest.yaml")


@app.route("/load_yaml", methods=["POST"])
def load_yaml_route():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file uploaded"}), 400
    try:
        data = yaml.safe_load(f.read())
        if not isinstance(data, dict):
            return jsonify({"error": "Invalid YAML"}), 400
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 400


if __name__ == "__main__":
    port = 5000
    url  = f"http://localhost:{port}"
    print(f"\n  Telco Threat Intel Report Generator")
    print(f"  ─────────────────────────────────────")
    print(f"  Open in browser: {url}")
    print(f"  Press Ctrl+C to stop\n")
    threading.Timer(1.2, lambda: webbrowser.open(url)).start()
    app.run(host="127.0.0.1", port=port, debug=False)
