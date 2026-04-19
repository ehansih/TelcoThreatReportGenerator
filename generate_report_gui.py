"""
Telco Threat Intelligence Report Generator — Windows GUI
Double-click this file or run: python generate_report_gui.py

Requirements: pip install reportlab pyyaml
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import os
import sys
import datetime
import yaml

# ── Try importing reportlab ────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table,
        TableStyle, PageBreak, HRFlowable
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.colors import HexColor
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# ── Try importing python-docx ─────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── Colour palette ─────────────────────────────────────────────────────────────
C = {
    "bg":       "#0D1B2A", "mid":    "#1E2D3D", "card":   "#2A3D52",
    "teal":     "#00C2CB", "yel":    "#F4D35E", "red":    "#E63946",
    "orange":   "#F7A400", "green":  "#2ECC71", "white":  "#FFFFFF",
    "light":    "#C8D6E5", "muted":  "#7F8C8D",
}
MARGIN = 1.5 * 28.35   # 1.5 cm in points

# ═══════════════════════════════════════════════════════════════════════════════
#  PDF ENGINE
# ═══════════════════════════════════════════════════════════════════════════════

def hc(h): return HexColor(h)

def S(name, **kw): return ParagraphStyle(name, **kw)

ST = {}   # populated in build_pdf()

def init_styles():
    global ST
    ST = {
        "title":   S("T",  fontName="Helvetica-Bold", fontSize=26, leading=32,
                      textColor=hc(C["white"]), spaceAfter=6),
        "sub":     S("Su", fontName="Helvetica", fontSize=12, leading=16,
                      textColor=hc(C["teal"]), spaceAfter=4),
        "date":    S("D",  fontName="Helvetica-Bold", fontSize=10,
                      textColor=hc(C["light"])),
        "sh":      S("SH", fontName="Helvetica-Bold", fontSize=16, leading=20,
                      textColor=hc(C["teal"]), spaceBefore=12, spaceAfter=7),
        "ssh":     S("SS", fontName="Helvetica-Bold", fontSize=12, leading=16,
                      textColor=hc(C["yel"]), spaceBefore=9, spaceAfter=4),
        "body":    S("B",  fontName="Helvetica", fontSize=9.5, leading=14,
                      textColor=hc(C["light"]), alignment=TA_JUSTIFY, spaceAfter=5),
        "bullet":  S("Bu", fontName="Helvetica", fontSize=9.5, leading=13,
                      textColor=hc(C["light"]),
                      leftIndent=14, firstLineIndent=-10, spaceBefore=2, spaceAfter=2),
        "label":   S("L",  fontName="Helvetica-Bold", fontSize=8,
                      textColor=hc(C["teal"])),
        "small":   S("Sm", fontName="Helvetica", fontSize=8, leading=11,
                      textColor=hc(C["muted"])),
        "mono":    S("Mo", fontName="Courier", fontSize=8.5, leading=12,
                      textColor=hc(C["light"])),
        "rec":     S("Re", fontName="Helvetica", fontSize=9,
                      textColor=hc(C["yel"]), leading=13),
    }

def sev_col(s):
    return hc({"CRITICAL": C["red"], "HIGH": C["orange"],
               "MEDIUM": C["teal"], "LOW": C["green"]}.get(str(s).upper(), C["light"]))

def base_ts():
    return TableStyle([
        ("BACKGROUND",    (0,0), (-1, 0), hc(C["mid"])),
        ("TEXTCOLOR",     (0,0), (-1, 0), hc(C["teal"])),
        ("FONTNAME",      (0,0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0,0), (-1,-1), 8),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [hc(C["card"]), hc(C["bg"])]),
        ("TEXTCOLOR",     (0,1), (-1,-1), hc(C["light"])),
        ("FONTNAME",      (0,1), (-1,-1), "Helvetica"),
        ("ALIGN",         (0,0), (-1,-1), "LEFT"),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",    (0,0), (-1,-1), 5),
        ("BOTTOMPADDING", (0,0), (-1,-1), 5),
        ("LEFTPADDING",   (0,0), (-1,-1), 6),
        ("BOX",           (0,0), (-1,-1), 1, hc(C["teal"])),
        ("INNERGRID",     (0,0), (-1,-1), 0.3, hc(C["mid"])),
    ])


class DarkCanvas(rl_canvas.Canvas):
    def __init__(self, filename, report_title="", period="", tlp="AMBER", org="", **kw):
        super().__init__(filename, **kw)
        self._rt, self._period, self._tlp, self._org = report_title, period, tlp, org
        self._paint()

    def _paint(self):
        w, h = A4
        self.saveState()
        self.setFillColor(hc(C["bg"])); self.rect(0, 0, w, h, fill=1, stroke=0)
        self.setFillColor(hc(C["teal"])); self.rect(0, h-8*mm, w, 8*mm, fill=1, stroke=0)
        self.setFillColor(hc(C["bg"])); self.setFont("Helvetica-Bold", 8)
        self.drawString(MARGIN, h-5.5*mm, f"{self._rt.upper()}  |  {self._period}")
        self.drawRightString(w-MARGIN, h-5.5*mm, f"TLP:{self._tlp}")
        self.setFillColor(hc(C["mid"])); self.rect(0, 0, w, 10*mm, fill=1, stroke=0)
        self.setFillColor(hc(C["teal"])); self.setFont("Helvetica", 7.5)
        self.drawString(MARGIN, 3.5*mm, f"Classification: TLP:{self._tlp} — Not for public distribution.")
        self.setFillColor(hc(C["muted"]))
        self.drawRightString(w-MARGIN, 3.5*mm, f"Page {self._pageNumber}  |  {self._org}")
        self.setFillColor(hc(C["mid"])); self.rect(0, 10*mm, 3*mm, h-18*mm, fill=1, stroke=0)
        self.restoreState()

    def showPage(self):
        super().showPage(); self._paint()

    def save(self):
        super().save()


def build_pdf(data: dict, output_path: str, progress_cb=None):
    """Core PDF builder. data = dict from the GUI form."""
    init_styles()

    def prog(msg):
        if progress_cb: progress_cb(msg)

    prog("Initialising document...")

    title  = data.get("title",  "Telco Cyber Threat Intelligence Report")
    period = data.get("period", "")
    org    = data.get("org",    "")
    author = data.get("author", "")
    tlp    = data.get("tlp",    "AMBER")

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=MARGIN+5*mm, rightMargin=MARGIN,
        topMargin=MARGIN+8*mm, bottomMargin=MARGIN+10*mm,
        title=title, author=author)

    story = []

    # ── COVER ──────────────────────────────────────────────────────────────────
    prog("Building cover page...")
    story.append(Spacer(1, 2.5*cm))
    for line in title.split("—"):
        story.append(Paragraph(line.strip().upper(), ST["title"]))
    story.append(Paragraph(data.get("subtitle", "Communications Service Providers"), ST["sub"]))
    story.append(Spacer(1, 3*mm))
    story.append(Paragraph(period, ST["date"]))
    story.append(Spacer(1, 8*mm))

    tlv = data.get("threat_level", "HIGH")
    tl_tbl = Table([[Paragraph("OVERALL THREAT LEVEL", ST["label"]),
                     Paragraph(tlv, S("TLV", fontName="Helvetica-Bold", fontSize=14,
                                       textColor=sev_col(tlv), alignment=TA_CENTER))]],
                   colWidths=[100*mm, 60*mm])
    tl_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), hc(C["card"])),
        ("BOX",           (0,0), (-1,-1), 1.5, hc(C["teal"])),
        ("LINEAFTER",     (0,0), (0,-1), 1, hc(C["teal"])),
        ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",    (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
        ("LEFTPADDING",   (0,0), (-1,-1), 10),
    ]))
    story.append(tl_tbl)
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(
        f"Prepared by: {author}  |  Organization: {org}  |  "
        f"Date: {data.get('report_date','')}  |  TLP:{tlp}", ST["small"]))
    story.append(PageBreak())

    # ── EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
    prog("Building executive summary...")
    story.append(Paragraph("EXECUTIVE SUMMARY", ST["sh"]))
    story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))

    exec_para = data.get("exec_summary", "").strip()
    if exec_para:
        story.append(Paragraph(exec_para, ST["body"]))

    top_risks = [r.strip() for r in data.get("top_risks", "").splitlines() if r.strip()]
    if top_risks:
        story.append(Paragraph("Top Risks:", ST["label"]))
        for r in top_risks:
            story.append(Paragraph(f"• {r}", ST["bullet"]))
        story.append(Spacer(1, 4*mm))

    # Key findings table from text input
    kf_text = data.get("key_findings", "").strip()
    kf_rows = []
    for line in kf_text.splitlines():
        parts = [p.strip() for p in line.split("|")]
        if len(parts) >= 2:
            kf_rows.append(parts[:3] if len(parts) >= 3 else parts + ["HIGH"])
    if kf_rows:
        table_data = [["KEY FINDING", "THREAT ACTOR / VECTOR", "SEVERITY"]] + kf_rows
        tbl = Table(table_data, colWidths=[88*mm, 57*mm, 22*mm])
        ts  = base_ts()
        for i, row in enumerate(kf_rows, 1):
            ts.add("TEXTCOLOR", (2, i), (2, i), sev_col(row[2] if len(row)>2 else "HIGH"))
            ts.add("FONTNAME",  (2, i), (2, i), "Helvetica-Bold")
        tbl.setStyle(ts)
        story.append(tbl)
    story.append(PageBreak())

    # ── THREAT LANDSCAPE ───────────────────────────────────────────────────────
    prog("Building threat landscape...")
    story.append(Paragraph("THREAT LANDSCAPE", ST["sh"]))
    story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=10))

    domains = [
        ("signaling",    "Telephony Signaling Networks (SS7 / Diameter / GTP)"),
        ("core_5g",      "Core Network & Radio Access Networks (5G / LTE)"),
        ("enterprise_it","Enterprise IT — BSS/OSS & Identity Systems"),
        ("fraud",        "Telecom Fraud — SIM Swap, IRSF & AI-Enabled Vishing"),
        ("supply_chain", "Third-Party & Supply Chain Risk"),
    ]
    for key, domain_title in domains:
        overview   = data.get(f"{key}_overview", "").strip()
        incidents  = data.get(f"{key}_incidents", "").strip()
        mitigations= data.get(f"{key}_mitigations", "").strip()
        if not (overview or incidents or mitigations):
            continue
        story.append(Paragraph(domain_title, ST["ssh"]))
        if overview:
            story.append(Paragraph(overview, ST["body"]))
        if incidents:
            inc_rows = []
            for line in incidents.splitlines():
                if "|" in line:
                    parts = [p.strip() for p in line.split("|")]
                    # format: Title | Date | Severity | Description
                    if len(parts) >= 3:
                        desc = parts[3] if len(parts) > 3 else parts[0]
                        title_inc = parts[0]
                        date_inc  = parts[1]
                        sev_inc   = parts[2]
                        inc_rows.append([f"{title_inc} — {desc}", date_inc, sev_inc])
                else:
                    # plain text line = description only
                    inc_rows.append([line.strip(), "", ""])
            if inc_rows:
                tbl_data = [["RECENT INCIDENTS", "DATE", "SEVERITY"]] + inc_rows
                tbl = Table(tbl_data, colWidths=[112*mm, 20*mm, 35*mm])
                ts  = base_ts()
                for i, row in enumerate(inc_rows, 1):
                    if row[2]:
                        ts.add("TEXTCOLOR", (2,i), (2,i), sev_col(row[2]))
                        ts.add("FONTNAME",  (2,i), (2,i), "Helvetica-Bold")
                tbl.setStyle(ts)
                story.append(tbl)
                story.append(Spacer(1, 4*mm))
        if mitigations:
            story.append(Paragraph(f"Mitigations — {domain_title}", ST["label"]))
            for m in mitigations.splitlines():
                if m.strip():
                    story.append(Paragraph(f"• {m.strip()}", ST["bullet"]))
        story.append(Spacer(1, 5*mm))

    story.append(PageBreak())

    # ── ACTORS & TTPs ──────────────────────────────────────────────────────────
    prog("Building actors & TTPs...")
    actors_text = data.get("actors", "").strip()
    if actors_text:
        story.append(Paragraph("THREAT ACTORS & TTPs", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        actor_rows = []
        for line in actors_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 6:
                    parts.append("")
                actor_rows.append(parts[:6])
        if actor_rows:
            cols = ["ACTOR", "ORIGIN", "TARGET", "INITIAL ACCESS", "KEY TTPs", "MITRE IDs"]
            tbl = Table([cols] + actor_rows,
                        colWidths=[22*mm, 18*mm, 25*mm, 30*mm, 47*mm, 25*mm])
            ts  = base_ts()
            state_kw = ["china","russia","iran","dprk","prc","state"]
            for i, row in enumerate(actor_rows, 1):
                origin = row[1].lower() if len(row) > 1 else ""
                col = hc(C["red"]) if any(k in origin for k in state_kw) else hc(C["orange"])
                ts.add("TEXTCOLOR", (1,i), (1,i), col)
                ts.add("FONTNAME",  (1,i), (1,i), "Helvetica-Bold")
                ts.add("TEXTCOLOR", (5,i), (5,i), hc(C["teal"]))
            tbl.setStyle(ts)
            story.append(tbl)
        story.append(PageBreak())

    # ── VULNERABILITIES ────────────────────────────────────────────────────────
    prog("Building vulnerabilities section...")
    vulns_text = data.get("vulnerabilities", "").strip()
    if vulns_text:
        story.append(Paragraph("VULNERABILITIES & EXPLOIT ALERT", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        vuln_rows = []
        for line in vulns_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 5: parts.append("")
                vuln_rows.append(parts[:5])
        if vuln_rows:
            cols = ["CVE ID", "VENDOR / PRODUCT", "DESCRIPTION", "CVSS", "PATCH PRIORITY"]
            tbl = Table([cols] + vuln_rows,
                        colWidths=[25*mm, 28*mm, 80*mm, 12*mm, 22*mm])
            ts  = base_ts()
            pri_map = {"IMMEDIATE": C["red"], "HIGH": C["orange"],
                       "MEDIUM": C["teal"], "LOW": C["green"]}
            for i, row in enumerate(vuln_rows, 1):
                pri = row[4].upper() if len(row) > 4 else ""
                if pri in pri_map:
                    ts.add("TEXTCOLOR", (4,i), (4,i), hc(pri_map[pri]))
                    ts.add("FONTNAME",  (4,i), (4,i), "Helvetica-Bold")
            tbl.setStyle(ts)
            story.append(tbl)
        story.append(Spacer(1, 6*mm))

    # ── EMERGING MALWARE ───────────────────────────────────────────────────────
    prog("Building malware section...")
    malware_text = data.get("malware", "").strip()
    if malware_text:
        story.append(Paragraph("EMERGING MALWARE NOTIFICATION", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        # Format: Name | Date | Severity | Description | Recommendation
        for line in malware_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 5: parts.append("")
                name, date, sev, desc, rec = parts[0], parts[1], parts[2], parts[3], parts[4]
                sc = sev_col(sev)
                hdr_tbl = Table([[
                    Paragraph(name, S("MH", fontName="Helvetica-Bold", fontSize=10,
                                       textColor=hc(C["yel"]))),
                    Paragraph(date, ST["small"]),
                    Paragraph(sev, S("MS", fontName="Helvetica-Bold", fontSize=9,
                                      textColor=sc, alignment=TA_CENTER)),
                ]], colWidths=[100*mm, 30*mm, 37*mm])
                hdr_tbl.setStyle(TableStyle([
                    ("BACKGROUND",    (0,0), (-1,-1), hc(C["mid"])),
                    ("VALIGN",        (0,0), (-1,-1), "MIDDLE"),
                    ("TOPPADDING",    (0,0), (-1,-1), 6),
                    ("BOTTOMPADDING", (0,0), (-1,-1), 6),
                    ("LEFTPADDING",   (0,0), (-1,-1), 8),
                    ("BOX",           (0,0), (-1,-1), 1, hc(C["teal"])),
                ]))
                story.append(hdr_tbl)
                if desc:
                    bt = Table([[Paragraph(f"<b>Description:</b> {desc}", ST["body"])]],
                                colWidths=[167*mm])
                    bt.setStyle(TableStyle([
                        ("BACKGROUND",    (0,0), (-1,-1), hc(C["card"])),
                        ("LEFTPADDING",   (0,0), (-1,-1), 8),
                        ("RIGHTPADDING",  (0,0), (-1,-1), 8),
                        ("TOPPADDING",    (0,0), (-1,-1), 5),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ]))
                    story.append(bt)
                if rec:
                    rt = Table([[Paragraph(f"<b>Recommendation:</b> {rec}", ST["rec"])]],
                                colWidths=[167*mm])
                    rt.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (-1,-1), HexColor("#1A2B1A")),
                        ("LEFTPADDING",   (0,0), (-1,-1), 8),
                        ("RIGHTPADDING",  (0,0), (-1,-1), 8),
                        ("TOPPADDING",    (0,0), (-1,-1), 5),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
                        ("BOX", (0,0), (-1,-1), 0.5, hc(C["green"])),
                    ]))
                    story.append(rt)
                story.append(Spacer(1, 4*mm))

    # ── BREACH NOTIFICATIONS ───────────────────────────────────────────────────
    prog("Building breach notifications...")
    breaches_text = data.get("breaches", "").strip()
    if breaches_text:
        story.append(Paragraph("BREACH NOTIFICATIONS", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        breach_rows = []
        for line in breaches_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 6: parts.append("")
                breach_rows.append(parts[:6])
        if breach_rows:
            cols = ["ORGANIZATION", "DATE", "SECTOR", "RECORDS EXPOSED", "VECTOR", "STATUS"]
            tbl = Table([cols] + breach_rows,
                        colWidths=[32*mm, 15*mm, 18*mm, 35*mm, 38*mm, 22*mm])
            ts  = base_ts()
            for i, row in enumerate(breach_rows, 1):
                status = row[5].lower() if len(row) > 5 else ""
                if "investigation" in status or "alleged" in status:
                    ts.add("TEXTCOLOR", (5,i), (5,i), hc(C["orange"]))
                    ts.add("FONTNAME",  (5,i), (5,i), "Helvetica-Bold")
            tbl.setStyle(ts)
            story.append(tbl)
        story.append(Spacer(1, 6*mm))

    # ── IOCs ──────────────────────────────────────────────────────────────────
    prog("Building IOC section...")
    iocs_text = data.get("iocs", "").strip()
    if iocs_text:
        story.append(Paragraph("INDICATORS OF COMPROMISE (IOCs)", ST["sh"]))
        story.append(HRFlowable(width="100%", thickness=1, color=hc(C["teal"]), spaceAfter=8))
        story.append(Paragraph(
            "Defanged IOCs for blocking/alerting. Full list in companion CSV file.", ST["small"]))
        story.append(Spacer(1, 3*mm))
        ioc_rows = []
        for line in iocs_text.splitlines():
            if "|" in line:
                parts = [p.strip() for p in line.split("|")]
                while len(parts) < 5: parts.append("")
                ioc_rows.append([
                    parts[0], Paragraph(parts[1], ST["mono"]),
                    parts[2], parts[3], parts[4]
                ])
        if ioc_rows:
            cols = ["TYPE", "VALUE", "CAMPAIGN / MALWARE", "CONFIDENCE", "ACTION"]
            tbl = Table([cols] + ioc_rows,
                        colWidths=[22*mm, 42*mm, 55*mm, 18*mm, 30*mm])
            ts  = base_ts()
            action_map = {"BLOCK": C["red"], "BLOCK AT STP": C["red"],
                          "ALERT": C["orange"], "MONITOR": C["teal"]}
            conf_map   = {"HIGH": C["green"], "MEDIUM": C["orange"], "LOW": C["muted"]}
            for i, row in enumerate(ioc_rows, 1):
                act  = str(row[4]).upper()
                conf = str(row[3]).upper()
                if act in action_map:
                    ts.add("TEXTCOLOR", (4,i), (4,i), hc(action_map[act]))
                    ts.add("FONTNAME",  (4,i), (4,i), "Helvetica-Bold")
                if conf in conf_map:
                    ts.add("TEXTCOLOR", (3,i), (3,i), hc(conf_map[conf]))
            tbl.setStyle(ts)
            story.append(tbl)

    # ── DISCLAIMER ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=hc(C["muted"]), spaceAfter=5))
    story.append(Paragraph(
        f"CLASSIFICATION: TLP:{tlp} — This report may not be shared beyond the recipient's "
        "organization. Based on information available at time of publication. Provided AS-IS "
        "without warranty. All IOCs are defanged for safe distribution.",
        ST["small"]))

    prog("Rendering PDF...")

    class BoundCanvas(DarkCanvas):
        def __init__(self, filename, **kw):
            super().__init__(filename,
                report_title=title, period=period, tlp=tlp, org=org, **kw)

    doc.build(story, canvasmaker=BoundCanvas)
    prog(f"Done! Saved to: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX ENGINE
# ═══════════════════════════════════════════════════════════════════════════════

def build_docx(data: dict, output_path: str, progress_cb=None):
    """Build a Word DOCX version of the threat intelligence report."""
    def prog(msg):
        if progress_cb: progress_cb(msg)

    prog("Initialising DOCX document...")

    title  = data.get("title",  "Telco Cyber Threat Intelligence Report")
    period = data.get("period", "")
    org    = data.get("org",    "")
    author = data.get("author", "")
    tlp    = data.get("tlp",    "AMBER")
    tlv    = data.get("threat_level", "HIGH")

    doc = DocxDocument()

    # ── Page margins ──────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def hex2rgb(h):
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def sev_rgb(s):
        return hex2rgb({"CRITICAL": "E63946", "HIGH": "F7A400",
                        "MEDIUM": "00C2CB", "LOW": "2ECC71"}.get(str(s).upper(), "C8D6E5"))

    def cell_bg(cell, hex_col):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_col.lstrip("#"))
        tcPr.append(shd)

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text.upper() if level == 1 else text)
        run.font.size  = Pt(16 if level == 1 else 11)
        run.font.bold  = True
        run.font.color.rgb = hex2rgb("0D1B2A" if level == 1 else "00C2CB")
        p.space_before = Pt(14 if level == 1 else 10)
        p.space_after  = Pt(6)
        if level == 1:
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "00C2CB")
            pBdr.append(bot); pPr.append(pBdr)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size  = Pt(9.5)
        p.space_after  = Pt(2)
        return p

    def make_table(headers, rows, col_count):
        tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
        tbl.style = "Table Grid"
        hrow = tbl.rows[0]
        for i, col in enumerate(headers):
            cell_bg(hrow.cells[i], "0D1B2A")
            run_c = hrow.cells[i].paragraphs[0].add_run(col)
            run_c.font.bold = True
            run_c.font.size = Pt(8)
            run_c.font.color.rgb = hex2rgb("00C2CB")
        return tbl

    # ── Header / Footer ───────────────────────────────────────────────────────
    section = doc.sections[0]
    hdr_para = section.header.paragraphs[0]
    hdr_para.text = f"{title}  |  {period}  |  TLP:{tlp}"
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hdr_para.runs:
        r.font.size = Pt(8); r.font.bold = True
        r.font.color.rgb = hex2rgb("0D1B2A")

    ftr_para = section.footer.paragraphs[0]
    ftr_para.text = f"Classification: TLP:{tlp} — Not for public distribution.  |  {org}"
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in ftr_para.runs:
        r.font.size = Pt(7)
        r.font.color.rgb = hex2rgb("7F8C8D")

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    prog("Building cover page...")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(72)
    run = p.add_run(title.upper())
    run.font.size = Pt(22); run.font.bold = True
    run.font.color.rgb = hex2rgb("0D1B2A")

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(data.get("subtitle", "Communications Service Providers"))
    run2.font.size = Pt(13); run2.font.color.rgb = hex2rgb("00C2CB")

    doc.add_paragraph()

    # Threat level table
    tl_tbl = doc.add_table(rows=1, cols=2)
    tl_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_bg(tl_tbl.cell(0, 0), "0D1B2A"); cell_bg(tl_tbl.cell(0, 1), "0D1B2A")
    r0 = tl_tbl.cell(0, 0).paragraphs[0].add_run("OVERALL THREAT LEVEL")
    r0.font.size = Pt(9); r0.font.bold = True; r0.font.color.rgb = hex2rgb("00C2CB")
    tl_tbl.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = tl_tbl.cell(0, 1).paragraphs[0].add_run(tlv)
    r1.font.size = Pt(16); r1.font.bold = True; r1.font.color.rgb = sev_rgb(tlv)

    doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Prepared by: {author}  |  Organization: {org}  |  "
                      f"Date: {data.get('report_date','')}  |  TLP:{tlp}")
    mr.font.size = Pt(9); mr.font.color.rgb = hex2rgb("7F8C8D")
    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────
    prog("Building executive summary...")
    add_heading("Executive Summary")

    exec_para = data.get("exec_summary", "").strip()
    if exec_para:
        ep = doc.add_paragraph(exec_para)
        if ep.runs: ep.runs[0].font.size = Pt(10)

    top_risks = [r.strip() for r in data.get("top_risks", "").splitlines() if r.strip()]
    if top_risks:
        lp = doc.add_paragraph()
        lr = lp.add_run("Top Risks:")
        lr.font.bold = True; lr.font.size = Pt(10)
        lr.font.color.rgb = hex2rgb("00C2CB")
        for r in top_risks:
            add_bullet(r)

    kf_text = data.get("key_findings", "").strip()
    kf_rows = []
    for line in kf_text.splitlines():
        parts = [pp.strip() for pp in line.split("|")]
        if len(parts) >= 2:
            kf_rows.append((parts + ["HIGH"])[:3])
    if kf_rows:
        doc.add_paragraph()
        lp = doc.add_paragraph()
        lr = lp.add_run("Key Findings"); lr.font.bold = True
        lr.font.size = Pt(10); lr.font.color.rgb = hex2rgb("00C2CB")
        tbl = make_table(["KEY FINDING", "ACTOR / VECTOR", "SEVERITY"], kf_rows, 3)
        for ri, row_data in enumerate(kf_rows, 1):
            cells = tbl.rows[ri].cells
            bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
            for ci, text in enumerate(row_data):
                cell_bg(cells[ci], bg); cells[ci].text = text
                if ci == 2 and text:
                    for para in cells[ci].paragraphs:
                        for run in para.runs:
                            run.font.bold = True; run.font.color.rgb = sev_rgb(text)

    doc.add_page_break()

    # ── THREAT LANDSCAPE ─────────────────────────────────────────────────────
    prog("Building threat landscape...")
    add_heading("Threat Landscape")

    domains = [
        ("signaling",    "Telephony Signaling Networks (SS7 / Diameter / GTP)"),
        ("core_5g",      "Core Network & Radio Access Networks (5G / LTE)"),
        ("enterprise_it","Enterprise IT — BSS/OSS & Identity Systems"),
        ("fraud",        "Telecom Fraud — SIM Swap, IRSF & AI-Enabled Vishing"),
        ("supply_chain", "Third-Party & Supply Chain Risk"),
    ]
    for key, domain_title in domains:
        overview    = data.get(f"{key}_overview", "").strip()
        incidents   = data.get(f"{key}_incidents", "").strip()
        mitigations = data.get(f"{key}_mitigations", "").strip()
        if not (overview or incidents or mitigations):
            continue
        add_heading(domain_title, level=2)
        if overview:
            ep = doc.add_paragraph(overview)
            if ep.runs: ep.runs[0].font.size = Pt(10)
        if incidents:
            inc_rows = []
            for line in incidents.splitlines():
                if "|" in line:
                    parts = [pp.strip() for pp in line.split("|")]
                    while len(parts) < 4: parts.append("")
                    inc_rows.append(parts[:4])
            if inc_rows:
                lp = doc.add_paragraph()
                lr = lp.add_run("Recent Incidents")
                lr.font.bold = True; lr.font.size = Pt(9)
                lr.font.color.rgb = hex2rgb("F7A400")
                tbl = make_table(["INCIDENT", "DATE", "SEVERITY", "DESCRIPTION"],
                                 inc_rows, 4)
                for ri, parts in enumerate(inc_rows, 1):
                    cells = tbl.rows[ri].cells
                    bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
                    for ci, text in enumerate(parts):
                        cell_bg(cells[ci], bg); cells[ci].text = text
                        if ci == 2 and text:
                            for para in cells[ci].paragraphs:
                                for run in para.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = sev_rgb(text)
                doc.add_paragraph()
        if mitigations:
            lp = doc.add_paragraph()
            lr = lp.add_run("Mitigations")
            lr.font.bold = True; lr.font.size = Pt(9)
            lr.font.color.rgb = hex2rgb("2ECC71")
            for m in mitigations.splitlines():
                if m.strip(): add_bullet(m.strip())

    doc.add_page_break()

    # ── THREAT ACTORS & TTPs ─────────────────────────────────────────────────
    actors_text = data.get("actors", "").strip()
    if actors_text:
        prog("Building actors & TTPs...")
        add_heading("Threat Actors & TTPs")
        actor_rows = []
        for line in actors_text.splitlines():
            if "|" in line:
                parts = [pp.strip() for pp in line.split("|")]
                while len(parts) < 6: parts.append("")
                actor_rows.append(parts[:6])
        if actor_rows:
            cols = ["ACTOR", "ORIGIN", "TARGET", "INITIAL ACCESS", "KEY TTPs", "MITRE IDs"]
            tbl = make_table(cols, actor_rows, 6)
            state_kw = ["china", "russia", "iran", "dprk", "prc", "state"]
            for ri, row_data in enumerate(actor_rows, 1):
                cells = tbl.rows[ri].cells
                bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
                for ci, text in enumerate(row_data):
                    cell_bg(cells[ci], bg); cells[ci].text = text
                    if ci == 1:
                        col_rgb = hex2rgb("E63946") if any(k in text.lower() for k in state_kw) \
                                  else hex2rgb("F7A400")
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.font.bold = True; run.font.color.rgb = col_rgb
                    if ci == 5:
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = hex2rgb("00C2CB")

    # ── VULNERABILITIES ───────────────────────────────────────────────────────
    vulns_text = data.get("vulnerabilities", "").strip()
    if vulns_text:
        prog("Building vulnerabilities...")
        add_heading("Vulnerabilities & Exploit Alert")
        vuln_rows = []
        for line in vulns_text.splitlines():
            if "|" in line:
                parts = [pp.strip() for pp in line.split("|")]
                while len(parts) < 5: parts.append("")
                vuln_rows.append(parts[:5])
        if vuln_rows:
            cols = ["CVE ID", "VENDOR / PRODUCT", "DESCRIPTION", "CVSS", "PATCH PRIORITY"]
            tbl = make_table(cols, vuln_rows, 5)
            pri_map = {"IMMEDIATE": "E63946", "HIGH": "F7A400",
                       "MEDIUM": "00C2CB", "LOW": "2ECC71"}
            for ri, row_data in enumerate(vuln_rows, 1):
                cells = tbl.rows[ri].cells
                bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
                for ci, text in enumerate(row_data):
                    cell_bg(cells[ci], bg); cells[ci].text = text
                    if ci == 4 and text.upper() in pri_map:
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = hex2rgb(pri_map[text.upper()])

    # ── EMERGING MALWARE ──────────────────────────────────────────────────────
    malware_text = data.get("malware", "").strip()
    if malware_text:
        prog("Building malware section...")
        add_heading("Emerging Malware Notification")
        for line in malware_text.splitlines():
            if "|" in line:
                parts = [pp.strip() for pp in line.split("|")]
                while len(parts) < 5: parts.append("")
                name, date, sev, desc, rec = parts[0], parts[1], parts[2], parts[3], parts[4]
                hp = doc.add_paragraph()
                hr = hp.add_run(f"[{sev}]  {name}  —  {date}")
                hr.font.bold = True; hr.font.size = Pt(11)
                hr.font.color.rgb = sev_rgb(sev)
                if desc:
                    dp = doc.add_paragraph()
                    dr = dp.add_run("Description: "); dr.font.bold = True; dr.font.size = Pt(9.5)
                    dr2 = dp.add_run(desc); dr2.font.size = Pt(9.5)
                if rec:
                    rp = doc.add_paragraph()
                    rr = rp.add_run("Recommendation: ")
                    rr.font.bold = True; rr.font.size = Pt(9.5)
                    rr.font.color.rgb = hex2rgb("2ECC71")
                    rr2 = rp.add_run(rec); rr2.font.size = Pt(9.5)
                doc.add_paragraph()

    # ── BREACH NOTIFICATIONS ──────────────────────────────────────────────────
    breaches_text = data.get("breaches", "").strip()
    if breaches_text:
        prog("Building breach notifications...")
        add_heading("Breach Notifications")
        breach_rows = []
        for line in breaches_text.splitlines():
            if "|" in line:
                parts = [pp.strip() for pp in line.split("|")]
                while len(parts) < 6: parts.append("")
                breach_rows.append(parts[:6])
        if breach_rows:
            cols = ["ORGANIZATION", "DATE", "SECTOR", "RECORDS EXPOSED", "VECTOR", "STATUS"]
            tbl = make_table(cols, breach_rows, 6)
            for ri, row_data in enumerate(breach_rows, 1):
                cells = tbl.rows[ri].cells
                bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
                for ci, text in enumerate(row_data):
                    cell_bg(cells[ci], bg); cells[ci].text = text
                    if ci == 5:
                        status = text.lower()
                        if "investigation" in status or "alleged" in status:
                            for para in cells[ci].paragraphs:
                                for run in para.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = hex2rgb("F7A400")

    # ── IOCs ──────────────────────────────────────────────────────────────────
    iocs_text = data.get("iocs", "").strip()
    if iocs_text:
        prog("Building IOC section...")
        add_heading("Indicators of Compromise (IOCs)")
        np = doc.add_paragraph(
            "Defanged IOCs for blocking/alerting. All IOCs are safe to distribute.")
        if np.runs: np.runs[0].font.size = Pt(9); np.runs[0].font.italic = True
        ioc_rows = []
        for line in iocs_text.splitlines():
            if "|" in line:
                parts = [pp.strip() for pp in line.split("|")]
                while len(parts) < 5: parts.append("")
                ioc_rows.append(parts[:5])
        if ioc_rows:
            cols = ["TYPE", "VALUE", "CAMPAIGN / MALWARE", "CONFIDENCE", "ACTION"]
            tbl = make_table(cols, ioc_rows, 5)
            action_map = {"BLOCK": "E63946", "BLOCK AT STP": "E63946",
                          "ALERT": "F7A400", "MONITOR": "00C2CB"}
            conf_map   = {"HIGH": "2ECC71", "MEDIUM": "F7A400", "LOW": "7F8C8D"}
            for ri, row_data in enumerate(ioc_rows, 1):
                cells = tbl.rows[ri].cells
                bg = "F5F7FA" if ri % 2 == 0 else "FFFFFF"
                for ci, text in enumerate(row_data):
                    cell_bg(cells[ci], bg); cells[ci].text = text
                    if ci == 1:
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.font.name = "Courier New"; run.font.size = Pt(8.5)
                    if ci == 3 and text.upper() in conf_map:
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = hex2rgb(conf_map[text.upper()])
                    if ci == 4 and text.upper() in action_map:
                        for para in cells[ci].paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = hex2rgb(action_map[text.upper()])

    # ── DISCLAIMER ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    dp = doc.add_paragraph()
    dr = dp.add_run(
        f"CLASSIFICATION: TLP:{tlp} — This report may not be shared beyond the recipient's "
        "organization. Based on information available at time of publication. Provided AS-IS "
        "without warranty. All IOCs are defanged for safe distribution.")
    dr.font.size = Pt(8); dr.font.italic = True
    dr.font.color.rgb = hex2rgb("7F8C8D")

    prog("Rendering DOCX...")
    doc.save(output_path)
    prog(f"Done! Saved to: {output_path}")
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  TKINTER GUI
# ═══════════════════════════════════════════════════════════════════════════════

DARK  = "#1a1a2e"
DARK2 = "#16213e"
DARK3 = "#0f3460"
TEAL  = "#00C2CB"
YEL   = "#F4D35E"
RED   = "#E63946"
WHITE = "#e0e0e0"
GREY  = "#7f8c8d"


class ScrollableFrame(ttk.Frame):
    def __init__(self, parent, **kw):
        super().__init__(parent, **kw)
        self.canvas = tk.Canvas(self, bg=DARK, highlightthickness=0)
        self.vsb    = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.inner  = ttk.Frame(self.canvas)
        self.canvas.configure(yscrollcommand=self.vsb.set)
        self.vsb.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.inner, anchor="nw")
        self.inner.bind("<Configure>", self._on_frame_configure)
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)

    def _on_frame_configure(self, e):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _on_canvas_configure(self, e):
        self.canvas.itemconfig(self.canvas_frame, width=e.width)

    def _on_mousewheel(self, e):
        self.canvas.yview_scroll(int(-1*(e.delta/120)), "units")


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Telco Threat Intelligence Report Generator")
        self.geometry("900x780")
        self.configure(bg=DARK)
        self.resizable(True, True)

        # Style
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(".",              background=DARK, foreground=WHITE,
                         fieldbackground=DARK2, font=("Segoe UI", 9))
        style.configure("TFrame",        background=DARK)
        style.configure("TLabel",        background=DARK, foreground=WHITE)
        style.configure("TLabelframe",   background=DARK, foreground=TEAL,
                         bordercolor=TEAL)
        style.configure("TLabelframe.Label", background=DARK, foreground=TEAL,
                         font=("Segoe UI", 9, "bold"))
        style.configure("TNotebook",     background=DARK2, bordercolor=DARK3)
        style.configure("TNotebook.Tab", background=DARK3, foreground=WHITE,
                         padding=[10, 4])
        style.map("TNotebook.Tab",
                  background=[("selected", TEAL)],
                  foreground=[("selected", DARK)])
        style.configure("TEntry",        fieldbackground=DARK2, foreground=WHITE,
                         bordercolor=DARK3, insertcolor=WHITE)
        style.configure("TCombobox",     fieldbackground=DARK2, foreground=WHITE,
                         selectbackground=DARK3)
        style.configure("Gen.TButton",   background=TEAL, foreground=DARK,
                         font=("Segoe UI", 10, "bold"), padding=8)
        style.map("Gen.TButton",
                  background=[("active", "#00a8b0")])
        style.configure("Docx.TButton", background="#2B579A", foreground=WHITE,
                         font=("Segoe UI", 10, "bold"), padding=8)
        style.map("Docx.TButton",
                  background=[("active", "#1e3f70")])
        style.configure("Sec.TButton",   background=DARK3, foreground=WHITE,
                         font=("Segoe UI", 9), padding=6)
        style.map("Sec.TButton",
                  background=[("active", DARK2)])

        self._build_header()
        self._build_notebook()
        self._build_footer()

    # ── HEADER ────────────────────────────────────────────────────────────────
    def _build_header(self):
        hdr = tk.Frame(self, bg=DARK3, height=60)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)
        tk.Label(hdr, text="TELCO THREAT INTEL REPORT GENERATOR",
                 bg=DARK3, fg=TEAL,
                 font=("Segoe UI", 14, "bold")).pack(side="left", padx=16, pady=10)
        tk.Label(hdr, text="v1.0  |  TLP:AMBER",
                 bg=DARK3, fg=YEL,
                 font=("Segoe UI", 9)).pack(side="right", padx=16)

    # ── NOTEBOOK TABS ─────────────────────────────────────────────────────────
    def _build_notebook(self):
        self.nb = ttk.Notebook(self)
        self.nb.pack(fill="both", expand=True, padx=8, pady=6)

        tabs = [
            ("Report Info",       self._tab_info),
            ("Executive Summary", self._tab_exec),
            ("Threat Landscape",  self._tab_landscape),
            ("Actors & CVEs",     self._tab_actors),
            ("Malware & Breaches",self._tab_malware),
            ("IOCs",              self._tab_iocs),
        ]
        self.frames = {}
        for name, builder in tabs:
            sf = ScrollableFrame(self.nb)
            self.nb.add(sf, text=f"  {name}  ")
            self.frames[name] = sf.inner
            builder(sf.inner)

    # ── TAB 1: REPORT INFO ────────────────────────────────────────────────────
    def _tab_info(self, parent):
        self.vars = {}

        def lf(p, text):
            f = ttk.LabelFrame(p, text=text, padding=10)
            f.pack(fill="x", padx=10, pady=5)
            return f

        def row(parent, label, var_key, default="", width=40, options=None):
            r = ttk.Frame(parent); r.pack(fill="x", pady=2)
            ttk.Label(r, text=label, width=22, anchor="w").pack(side="left")
            if options:
                self.vars[var_key] = tk.StringVar(value=default)
                cb = ttk.Combobox(r, textvariable=self.vars[var_key],
                                  values=options, width=width-5, state="readonly")
                cb.pack(side="left", fill="x", expand=True)
            else:
                self.vars[var_key] = tk.StringVar(value=default)
                ttk.Entry(r, textvariable=self.vars[var_key], width=width).pack(
                    side="left", fill="x", expand=True)

        f1 = lf(parent, "Report Metadata")
        row(f1, "Report Title",     "title",
            "Telco Cyber Threat Intelligence Report")
        row(f1, "Subtitle",         "subtitle",
            "Communications Service Providers — India & Global")
        row(f1, "Period Covered",   "period",  "April 2026")
        row(f1, "Report Date",      "report_date",
            datetime.date.today().strftime("%Y-%m-%d"))
        row(f1, "Author",           "author",  "")
        row(f1, "Organization",     "org",     "")

        f2 = lf(parent, "Classification")
        row(f2, "TLP Level", "tlp", "AMBER",
            options=["WHITE", "GREEN", "AMBER", "RED"])
        row(f2, "Overall Threat Level", "threat_level", "HIGH",
            options=["CRITICAL", "HIGH", "MEDIUM", "LOW"])

        f3 = lf(parent, "Output")
        # PDF path
        r = ttk.Frame(f3); r.pack(fill="x", pady=2)
        ttk.Label(r, text="Output PDF Path", width=22, anchor="w").pack(side="left")
        self.vars["output"] = tk.StringVar(
            value=os.path.join(os.path.expanduser("~"), "Downloads",
                               "Telco_TI_Report.pdf"))
        ttk.Entry(r, textvariable=self.vars["output"]).pack(side="left", fill="x", expand=True)
        ttk.Button(r, text="Browse", style="Sec.TButton",
                   command=self._browse_output).pack(side="left", padx=4)
        # DOCX path
        r2 = ttk.Frame(f3); r2.pack(fill="x", pady=2)
        ttk.Label(r2, text="Output DOCX Path", width=22, anchor="w").pack(side="left")
        self.vars["output_docx"] = tk.StringVar(
            value=os.path.join(os.path.expanduser("~"), "Downloads",
                               "Telco_TI_Report.docx"))
        ttk.Entry(r2, textvariable=self.vars["output_docx"]).pack(
            side="left", fill="x", expand=True)
        ttk.Button(r2, text="Browse", style="Sec.TButton",
                   command=self._browse_docx).pack(side="left", padx=4)

    def _browse_output(self):
        p = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf")],
            initialfile="Telco_TI_Report.pdf",
            initialdir=os.path.expanduser("~"))
        if p:
            self.vars["output"].set(p)

    def _browse_docx(self):
        p = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx")],
            initialfile="Telco_TI_Report.docx",
            initialdir=os.path.expanduser("~"))
        if p:
            self.vars["output_docx"].set(p)

    # ── TAB 2: EXECUTIVE SUMMARY ──────────────────────────────────────────────
    def _tab_exec(self, parent):
        self._labeled_text(parent, "Executive Summary Paragraph",
                           "exec_summary", height=5,
                           hint="Write 3-5 sentences summarising the overall threat assessment.")
        self._labeled_text(parent, "Top Risks (one per line)",
                           "top_risks", height=5,
                           hint="Each line = one risk bullet point.")
        self._labeled_text(parent,
            "Key Findings Table  (pipe-separated: Finding | Actor/Vector | Severity)",
            "key_findings", height=7,
            hint="Example:\nSS7 OTP interception active | Unattributed eCrime | CRITICAL\n"
                 "Salt Typhoon in backbone | China-nexus APT | HIGH")

    # ── TAB 3: THREAT LANDSCAPE ───────────────────────────────────────────────
    def _tab_landscape(self, parent):
        domains = [
            ("signaling",     "SS7 / Diameter / GTP Signaling"),
            ("core_5g",       "Core Network & 5G / LTE"),
            ("enterprise_it", "Enterprise IT (BSS/OSS)"),
            ("fraud",         "Telecom Fraud (SIM Swap / IRSF / Vishing)"),
            ("supply_chain",  "Supply Chain & Third-Party"),
        ]
        for key, label in domains:
            lf = ttk.LabelFrame(parent, text=label, padding=8)
            lf.pack(fill="x", padx=10, pady=5)
            self._inline_text(lf, f"{key}_overview",
                              "Overview (2-3 paragraphs):", height=3)
            self._inline_text(lf, f"{key}_incidents",
                              "Incidents  (Title | Date | Severity | Description):",
                              height=3,
                              hint="Example: SK Telecom BPFDoor | Apr 2025 | CRITICAL | 27M records stolen")
            self._inline_text(lf, f"{key}_mitigations",
                              "Mitigations (one per line):", height=3)

    # ── TAB 4: ACTORS & CVEs ──────────────────────────────────────────────────
    def _tab_actors(self, parent):
        self._labeled_text(parent,
            "Threat Actors  (pipe-separated: Name | Origin | Target | Initial Access | TTPs | MITRE IDs)",
            "actors", height=8,
            hint="Example:\nSalt Typhoon | China (PRC) | Tier-1 backbone | Exploit IOS-XE | LotL, packet capture | T1190, T1133\n"
                 "ShinyHunters | eCrime | CRM systems | Vishing + OAuth | PII exfil | T1566, T1119")

        self._labeled_text(parent,
            "Vulnerabilities  (pipe-separated: CVE ID | Vendor/Product | Description | CVSS | Patch Priority)",
            "vulnerabilities", height=8,
            hint="Example:\nCVE-2026-2329 | Grandstream GXP1600 | Unauthenticated RCE via web API | 9.8 | IMMEDIATE\n"
                 "CVE-2025-24200 | Apple iOS | USB Restricted Mode bypass | 6.1 | HIGH\n"
                 "Patch Priority options: IMMEDIATE | HIGH | MEDIUM | LOW")

    # ── TAB 5: MALWARE & BREACHES ─────────────────────────────────────────────
    def _tab_malware(self, parent):
        self._labeled_text(parent,
            "Emerging Malware  (pipe-separated: Name | Date | Severity | Description | Recommendation)",
            "malware", height=8,
            hint="Example:\nBPFDoor | Mar 2026 | HIGH | Passive BPF backdoor targeting Linux BSS/OSS | Audit raw socket listeners with 'ss -lp'\n"
                 "Severity options: CRITICAL | HIGH | MEDIUM")

        self._labeled_text(parent,
            "Breach Notifications  (pipe-separated: Organization | Date | Sector | Records Exposed | Vector | Status)",
            "breaches", height=8,
            hint="Example:\nSK Telecom | Apr 2025 | Telco | 27M subscriber records | BPFDoor/LotL | Confirmed\n"
                 "Status options: Confirmed | Under Investigation | Alleged | Contained")

    # ── TAB 6: IOCs ───────────────────────────────────────────────────────────
    def _tab_iocs(self, parent):
        note = tk.Label(parent,
            text="Always defang IOCs:  replace  .  with  [.]  in IPs and domains\n"
                 "Example:  185.220.101[.]47   or   malware-domain[.]com",
            bg=DARK2, fg=YEL, font=("Segoe UI", 9), justify="left", pady=6, padx=10)
        note.pack(fill="x", padx=10, pady=4)

        self._labeled_text(parent,
            "IOCs  (pipe-separated: Type | Value | Campaign/Malware | Confidence | Action)",
            "iocs", height=15,
            hint="Example:\nIPv4 | 185.220.101[.]47 | GhostSignal SS7 C2 | HIGH | BLOCK\n"
                 "Domain | update-secure[.]xyz | OysterLoader | HIGH | BLOCK\n"
                 "SCCP GT | +442012345XXX | Rogue HLR | HIGH | BLOCK at STP\n"
                 "File Hash (SHA256) | 3a9f2c1d...e84b7f92 | BPFDoor ELF | HIGH | BLOCK\n"
                 "Confidence: HIGH | MEDIUM | LOW     Action: BLOCK | ALERT | MONITOR | BLOCK at STP")

    # ── FOOTER ────────────────────────────────────────────────────────────────
    def _build_footer(self):
        foot = tk.Frame(self, bg=DARK2, pady=8)
        foot.pack(fill="x", side="bottom")

        self.progress_var = tk.StringVar(value="Ready")
        tk.Label(foot, textvariable=self.progress_var,
                 bg=DARK2, fg=GREY,
                 font=("Segoe UI", 8)).pack(side="left", padx=12)

        ttk.Button(foot, text="  GENERATE PDF  ",
                   style="Gen.TButton",
                   command=self._generate).pack(side="right", padx=4)

        ttk.Button(foot, text="  GENERATE DOCX  ",
                   style="Docx.TButton",
                   command=self._generate_docx).pack(side="right", padx=4)

        ttk.Button(foot, text="Load YAML Manifest",
                   style="Sec.TButton",
                   command=self._load_yaml).pack(side="right", padx=4)

        ttk.Button(foot, text="Save YAML Manifest",
                   style="Sec.TButton",
                   command=self._save_yaml).pack(side="right", padx=4)

    # ── HELPERS ───────────────────────────────────────────────────────────────
    def _labeled_text(self, parent, label, key, height=4, hint=""):
        if not hasattr(self, "texts"):
            self.texts = {}
        lf = ttk.LabelFrame(parent, text=label, padding=8)
        lf.pack(fill="x", padx=10, pady=5)
        if hint:
            tk.Label(lf, text=hint, bg=DARK, fg=GREY,
                     font=("Segoe UI", 8), justify="left",
                     wraplength=780).pack(anchor="w")
        txt = scrolledtext.ScrolledText(lf, height=height,
            bg=DARK2, fg=WHITE, insertbackground=WHITE,
            font=("Consolas", 9), relief="flat",
            wrap=tk.WORD)
        txt.pack(fill="x", pady=2)
        self.texts[key] = txt

    def _inline_text(self, parent, key, label, height=3, hint=""):
        if not hasattr(self, "texts"):
            self.texts = {}
        ttk.Label(parent, text=label, foreground=GREY).pack(anchor="w")
        if hint:
            tk.Label(parent, text=hint, bg=DARK, fg=GREY,
                     font=("Segoe UI", 7), justify="left").pack(anchor="w")
        txt = scrolledtext.ScrolledText(parent, height=height,
            bg=DARK2, fg=WHITE, insertbackground=WHITE,
            font=("Consolas", 9), relief="flat", wrap=tk.WORD)
        txt.pack(fill="x", pady=2)
        self.texts[key] = txt

    def _get_texts(self):
        return {k: v.get("1.0", "end-1c") for k, v in self.texts.items()}

    def _get_all_data(self):
        data = {k: v.get() for k, v in self.vars.items()}
        data.update(self._get_texts())
        return data

    # ── YAML LOAD / SAVE ─────────────────────────────────────────────────────
    def _save_yaml(self):
        path = filedialog.asksaveasfilename(
            defaultextension=".yaml",
            filetypes=[("YAML files", "*.yaml *.yml")],
            initialfile="my_report.yaml")
        if not path:
            return
        data = self._get_all_data()
        with open(path, "w") as f:
            yaml.dump(data, f, default_flow_style=False, allow_unicode=True)
        messagebox.showinfo("Saved", f"Manifest saved to:\n{path}")

    def _load_yaml(self):
        path = filedialog.askopenfilename(
            filetypes=[("YAML files", "*.yaml *.yml"), ("All files", "*.*")])
        if not path:
            return
        with open(path) as f:
            data = yaml.safe_load(f)
        if not isinstance(data, dict):
            messagebox.showerror("Error", "Invalid YAML file.")
            return
        for k, v in data.items():
            if k in self.vars:
                self.vars[k].set(str(v) if v else "")
            elif k in self.texts:
                self.texts[k].delete("1.0", "end")
                self.texts[k].insert("1.0", str(v) if v else "")
        messagebox.showinfo("Loaded", "Manifest loaded successfully.")

    # ── GENERATE ─────────────────────────────────────────────────────────────
    def _generate(self):
        if not REPORTLAB_OK:
            messagebox.showerror("Missing Dependency",
                "reportlab is not installed.\n\n"
                "Run in Command Prompt:\n  pip install reportlab pyyaml")
            return

        data   = self._get_all_data()
        output = data.get("output", "").strip()
        if not output:
            messagebox.showerror("Error", "Please set an output PDF path.")
            return

        self.progress_var.set("Generating...")

        def run():
            try:
                build_pdf(data, output, progress_cb=lambda m: self.progress_var.set(m))
                self.after(0, lambda: messagebox.showinfo(
                    "Done!", f"Report saved to:\n{output}"))
            except Exception as ex:
                self.after(0, lambda: messagebox.showerror(
                    "Error", f"Failed to generate PDF:\n\n{ex}"))
            finally:
                self.after(0, lambda: self.progress_var.set("Ready"))

        threading.Thread(target=run, daemon=True).start()

    def _generate_docx(self):
        if not DOCX_OK:
            messagebox.showerror("Missing Dependency",
                "python-docx is not installed.\n\n"
                "Run in Command Prompt:\n  pip install python-docx\n\n"
                "Then relaunch this application.")
            return

        data   = self._get_all_data()
        output = data.get("output_docx", "").strip()
        if not output:
            messagebox.showerror("Error", "Please set an output DOCX path.")
            return

        self.progress_var.set("Generating DOCX...")

        def run():
            try:
                build_docx(data, output, progress_cb=lambda m: self.progress_var.set(m))
                self.after(0, lambda: messagebox.showinfo(
                    "Done!", f"Word document saved to:\n{output}"))
            except Exception as ex:
                self.after(0, lambda: messagebox.showerror(
                    "Error", f"Failed to generate DOCX:\n\n{ex}"))
            finally:
                self.after(0, lambda: self.progress_var.set("Ready"))

        threading.Thread(target=run, daemon=True).start()


# ── ENTRY POINT ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if not REPORTLAB_OK:
        root = tk.Tk(); root.withdraw()
        messagebox.showerror("Missing Dependency",
            "reportlab is not installed.\n\n"
            "Open Command Prompt and run:\n"
            "  pip install reportlab pyyaml\n\n"
            "Then relaunch this application.")
        sys.exit(1)
    app = App()
    app.mainloop()
